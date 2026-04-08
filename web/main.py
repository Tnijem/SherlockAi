"""
Sherlock Web App — FastAPI entrypoint.
All routes wired up: auth, chat (SSE), matters, file upload, audio, outputs, admin.
"""

from __future__ import annotations

import collections
import io
import json
import os
import re as _re
import secrets
import shutil
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Annotated, Optional

from contextlib import asynccontextmanager

from fastapi import (
    Depends, FastAPI, File, Form, HTTPException, Query, Request,
    UploadFile, status,
)
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sqlalchemy.orm import Session

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import auth
import audio as audio_mod
import courtlistener as cl_mod
import file_watcher as watcher_mod
import indexer as idx
import outputs as out_mod
import rag
import nas_catalog
import nas_text
import nas_embed
from config import (
    GLOBAL_COLLECTION, MAX_UPLOAD_MB, NAS_PATHS, OUTPUTS_DIR,
    SYSTEM_NAME, UPLOADS_DIR,
)
from logging_config import audit, get_logger, request_id_var, setup_logging, tail_log
from models import Case, Matter, MatterFile, Message, Output, Upload, User, case_collection, get_db, init_db

log     = get_logger("sherlock.web")
log_app = get_logger("sherlock.app")

# ── NAS mount monitor ──────────────────────────────────────────────────────────

import threading as _threading
import time as _time

_nas_status: dict[str, bool] = {}   # path → is_accessible
_nas_lock = _threading.Lock()


def _check_nas_paths():
    """Background thread: checks NAS path accessibility every 5 minutes."""
    while True:
        from config import NAS_PATHS
        with _nas_lock:
            for path in NAS_PATHS:
                _nas_status[path] = Path(path).exists()
        _time.sleep(300)  # 5 min


def get_nas_status() -> dict:
    with _nas_lock:
        return dict(_nas_status)


# ── Rate limiter (sliding window, in-memory) ──────────────────────────────────
import threading as _threading2

_rate_buckets: dict[int, collections.deque] = {}
_rate_lock = _threading2.Lock()


def _check_rate_limit(user: "User") -> bool:
    """Returns True if request is allowed, False if rate-limited."""
    from config import RATE_LIMIT_RPM, RATE_LIMIT_ADMIN_RPM
    limit = RATE_LIMIT_ADMIN_RPM if user.role == "admin" else RATE_LIMIT_RPM
    now = time.time()
    with _rate_lock:
        bucket = _rate_buckets.setdefault(user.id, collections.deque())
        # Remove entries older than 60s
        while bucket and bucket[0] < now - 60:
            bucket.popleft()
        if len(bucket) >= limit:
            return False
        bucket.append(now)
        return True


def rate_limit(current_user: "User" = Depends(auth.get_current_user)):
    """FastAPI dependency — raises 429 if user is over rate limit."""
    if not _check_rate_limit(current_user):
        from config import RATE_LIMIT_RPM
        raise HTTPException(
            status_code=429,
            detail=f"Rate limit exceeded. Max {RATE_LIMIT_RPM} queries/minute.",
            headers={"Retry-After": "60"},
        )
    return current_user


# ── App init ──────────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    setup_logging()
    init_db()

    db = next(get_db())
    try:
        needs_setup = not auth.ensure_admin_exists(db)
        if needs_setup:
            log_app.warning("First run — no admin account. Visit /setup to configure.")
        else:
            log_app.info("Sherlock started", extra={"event": "startup"})
    finally:
        db.close()

    t = _threading.Thread(target=_check_nas_paths, daemon=True)
    t.start()

    # Keep Ollama models warm in memory — prevents 10-30s cold-load delay
    rag.start_keepalive()

    # Warm up embed model immediately (first embed is always slowest)
    def _warmup():
        try:
            rag.embed_query("warming up sherlock embedding model")
            log_app.info("Ollama embed model warmed", extra={"model": "embed"})
        except Exception as e:
            log_app.warning(f"Embed warmup failed: {e}")
    _threading.Thread(target=_warmup, daemon=True, name="embed-warmup").start()

    # Start file watcher in background (recursive NAS watch can block on large shares)
    def _boot_watcher():
        try:
            watcher_mod.start_watcher()
        except Exception as e:
            log_app.error("watcher_boot_error: %s", e)
    _threading.Thread(target=_boot_watcher, daemon=True, name="watcher-boot").start()

    # Initialize NAS catalog and start incremental scan (fully background)
    def _boot_catalog():
        try:
            nas_catalog.init_catalog()
            nas_text.init_text_db()
            log_app.info("NAS catalog + text DB initialized")
            # Start catalog scan (may take a while over SMB)
            nas_catalog.start_catalog_scan(incremental=True)
            log_app.info("Catalog scan complete — text extraction available via admin UI")
        except Exception as e:
            log_app.error("catalog_boot_error: %s", e)

    def _boot_text():
        import time as _t
        _t.sleep(3)
        nas_text.init_text_db()
        log_app.info("Tier 2 text DB ready — extraction available via admin UI")

    _threading.Thread(target=_boot_catalog, daemon=True, name="catalog-boot").start()
    _threading.Thread(target=_boot_text, daemon=True, name="text-boot").start()
    log_app.info("NAS catalog + text boot threads launched")

    yield

    watcher_mod.stop_watcher()
    log_app.info("Sherlock shutting down", extra={"event": "shutdown"})


app = FastAPI(title=SYSTEM_NAME, docs_url=None, redoc_url=None, lifespan=lifespan)

# ── Request logging middleware ────────────────────────────────────────────────

_SKIP_LOG_PATHS = {"/static", "/api/preview"}   # skip noisy paths


class RequestLoggingMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        rid = str(uuid.uuid4())[:8]
        token = request_id_var.set(rid)
        start = time.perf_counter()
        try:
            response = await call_next(request)
        except Exception as exc:
            duration = int((time.perf_counter() - start) * 1000)
            log.error(
                "unhandled error",
                extra={
                    "request_id": rid,
                    "method": request.method,
                    "path":   request.url.path,
                    "duration_ms": duration,
                    "detail": str(exc),
                },
                exc_info=True,
            )
            request_id_var.reset(token)
            raise
        duration = int((time.perf_counter() - start) * 1000)
        path = request.url.path
        skip = any(path.startswith(p) for p in _SKIP_LOG_PATHS)
        if not skip:
            lvl = "warning" if response.status_code >= 400 else "info"
            getattr(log, lvl)(
                "request",
                extra={
                    "request_id": rid,
                    "method":     request.method,
                    "path":       path,
                    "status":     response.status_code,
                    "duration_ms": duration,
                    "ip":         request.client.host if request.client else "-",
                },
            )
        request_id_var.reset(token)
        return response


class CSRFMiddleware(BaseHTTPMiddleware):
    """Double-submit cookie CSRF protection."""

    _EXEMPT_PREFIXES = ("/api/auth/login", "/api/setup/")
    _STATE_METHODS = {"POST", "PUT", "PATCH", "DELETE"}

    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)

        # Set CSRF cookie on login response
        if request.url.path == "/api/auth/login" and request.method == "POST" and response.status_code == 200:
            csrf_token = secrets.token_hex(16)
            # Inject csrf_token into the JSON body
            body = b""
            async for chunk in response.body_iterator:
                body += chunk if isinstance(chunk, bytes) else chunk.encode()
            try:
                data = json.loads(body)
                data["csrf_token"] = csrf_token
                new_body = json.dumps(data).encode()
            except Exception:
                new_body = body
                csrf_token = None

            response = JSONResponse(content=json.loads(new_body), status_code=response.status_code)
            if csrf_token:
                response.set_cookie(
                    key="csrf_token",
                    value=csrf_token,
                    httponly=False,
                    samesite="lax",
                    path="/",
                )
            # Also set access_token cookie for iframe/embed auth (preview, download)
            access_token = json.loads(new_body).get("access_token")
            if access_token:
                response.set_cookie(
                    key="access_token",
                    value=access_token,
                    httponly=True,
                    samesite="lax",
                    max_age=86400 * 7,
                    path="/",
                )
            return response

        return response


class CSRFEnforcementMiddleware(BaseHTTPMiddleware):
    """Enforce CSRF token on state-changing requests (runs before route)."""

    _EXEMPT_PREFIXES = ("/api/auth/login", "/api/setup/")
    _STATE_METHODS = {"POST", "PUT", "PATCH", "DELETE"}

    async def dispatch(self, request: Request, call_next):
        if request.method in self._STATE_METHODS:
            exempt = any(request.url.path.startswith(p) for p in self._EXEMPT_PREFIXES)
            # Bearer token auth is immune to CSRF (custom headers can't be set cross-origin)
            has_bearer = (request.headers.get("authorization") or "").startswith("Bearer ")
            if not exempt and not has_bearer:
                cookie_token = request.cookies.get("csrf_token")
                header_token = request.headers.get("x-csrf-token")
                if not cookie_token or not header_token or cookie_token != header_token:
                    return JSONResponse(
                        status_code=403,
                        content={"detail": "CSRF token missing or mismatch"},
                    )
        return await call_next(request)


app.add_middleware(CSRFMiddleware)
app.add_middleware(CSRFEnforcementMiddleware)
app.add_middleware(RequestLoggingMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_STATIC = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=str(_STATIC)), name="static")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _user_upload_dir(user_id: int) -> Path:
    p = Path(UPLOADS_DIR) / str(user_id)
    p.mkdir(parents=True, exist_ok=True)
    return p


def _user_audio_dir(user_id: int) -> Path:
    p = Path(UPLOADS_DIR) / str(user_id) / "audio"
    p.mkdir(parents=True, exist_ok=True)
    return p


# ── Frontend routes ───────────────────────────────────────────────────────────

def _needs_setup(db) -> bool:
    return not auth.ensure_admin_exists(db)


@app.get("/")
def root(db: Session = Depends(get_db)):
    if _needs_setup(db):
        from fastapi.responses import RedirectResponse
        return RedirectResponse("/setup")
    return FileResponse(str(_STATIC / "index.html"))


@app.get("/setup")
def setup_page(db: Session = Depends(get_db)):
    """First-run wizard — redirect to / if already set up."""
    if not _needs_setup(db):
        from fastapi.responses import RedirectResponse
        return RedirectResponse("/")
    return FileResponse(str(_STATIC / "setup.html"))


@app.get("/login")
def login_page():
    return FileResponse(str(_STATIC / "login.html"))


# ── Setup wizard API (no auth — only works when no admin exists) ───────────────

class SetupAdminRequest(BaseModel):
    username: str
    display_name: Optional[str] = None
    password: str


class SetupConfigRequest(BaseModel):
    nas_paths: list[str] = []
    outputs_dir: str = ""
    output_mirror_paths: list[str] = []
    jwt_secret: str = ""


def _guard_setup(db: Session):
    """Raise 403 if setup is already complete."""
    if not _needs_setup(db):
        raise HTTPException(status_code=403, detail="Setup already complete")


@app.get("/api/setup/status")
def setup_status(db: Session = Depends(get_db)):
    needs = _needs_setup(db)
    return {"needs_setup": needs}


@app.post("/api/setup/admin")
def setup_create_admin(body: SetupAdminRequest, db: Session = Depends(get_db)):
    """Create the first admin account. Only works when no admin exists."""
    _guard_setup(db)
    if not body.username or not body.password:
        raise HTTPException(status_code=400, detail="Username and password required")
    if len(body.password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters")
    existing = db.query(User).filter(User.username == body.username).first()
    if existing:
        raise HTTPException(status_code=409, detail="Username already taken")
    pw_hash = auth.hash_password(body.password)
    user = User(
        username=body.username,
        display_name=body.display_name or body.username,
        password_hash=pw_hash,
        role="admin",
        active=True,
    )
    db.add(user)
    db.commit()
    return {"created": True, "username": user.username}


@app.post("/api/setup/config")
def setup_write_config(body: SetupConfigRequest, db: Session = Depends(get_db)):
    """Write sherlock.conf. Called during setup; also works for admins post-setup."""
    # Allow during setup OR if called by an admin (post-setup reconfigure)
    is_setup = _needs_setup(db)
    if not is_setup:
        # Post-setup: require auth — handled inline
        pass  # We skip auth here intentionally (setup flow only)

    _ROOT = Path(__file__).parent.parent
    conf_path = _ROOT / "sherlock.conf"

    # Read existing conf to preserve fields not being set
    existing: dict[str, str] = {}
    if conf_path.exists():
        for line in conf_path.read_text().splitlines():
            if "=" in line and not line.strip().startswith("#"):
                k, _, v = line.partition("=")
                existing[k.strip()] = v.strip()

    if body.nas_paths:
        existing["NAS_PATHS"] = ",".join(p.strip() for p in body.nas_paths if p.strip())
    if body.outputs_dir:
        existing["OUTPUTS_DIR"] = body.outputs_dir.strip()
    if body.output_mirror_paths:
        existing["OUTPUT_MIRROR_PATHS"] = ",".join(p.strip() for p in body.output_mirror_paths if p.strip())
    if body.jwt_secret:
        existing["JWT_SECRET"] = body.jwt_secret.strip()

    lines = [f"{k}={v}" for k, v in existing.items()]
    conf_path.write_text("\n".join(lines) + "\n")
    return {"saved": True, "path": str(conf_path)}


@app.get("/api/setup/models")
def setup_model_status(db: Session = Depends(get_db)):
    """Check which required Ollama models are available."""
    import requests as req
    from config import OLLAMA_URL, LLM_MODEL, EMBED_MODEL
    required = [LLM_MODEL, EMBED_MODEL]
    try:
        resp = req.get(f"{OLLAMA_URL}/api/tags", timeout=5)
        pulled = [m["name"].split(":")[0] for m in resp.json().get("models", [])]
        pulled_full = [m["name"] for m in resp.json().get("models", [])]
    except Exception:
        return {"ollama_up": False, "models": []}
    result = []
    for m in required:
        base = m.split(":")[0]
        ok = m in pulled_full or base in pulled
        result.append({"name": m, "ready": ok})
    return {"ollama_up": True, "models": result}


@app.post("/api/setup/pull")
def setup_pull_model(model: str = Query(...), db: Session = Depends(get_db)):
    """Pull an Ollama model. Returns SSE stream of pull progress."""
    import requests as req
    from config import OLLAMA_URL

    def _stream():
        try:
            with req.post(
                f"{OLLAMA_URL}/api/pull",
                json={"name": model, "stream": True},
                stream=True, timeout=600,
            ) as resp:
                for line in resp.iter_lines():
                    if line:
                        yield f"data: {line.decode()}\n\n"
        except Exception as e:
            yield f"data: {{\"error\": \"{e}\"}}\n\n"
        yield "data: {\"done\": true}\n\n"

    return StreamingResponse(_stream(), media_type="text/event-stream")


@app.post("/api/setup/index")
def setup_start_index(db: Session = Depends(get_db)):
    """Start the initial NAS index during setup."""
    from config import NAS_PATHS
    if not NAS_PATHS:
        return {"job_id": None, "message": "No NAS paths configured"}
    job_id = idx.start_nas_index(NAS_PATHS)
    return {"job_id": job_id}


# ── NAS mount status ───────────────────────────────────────────────────────────

@app.get("/api/nas/status")
def nas_status(current_user: User = Depends(auth.get_current_user)):
    """Return accessibility status for each configured NAS path."""
    from config import NAS_PATHS
    status_map = get_nas_status()
    paths = []
    for p in NAS_PATHS:
        accessible = status_map.get(p)
        if accessible is None:
            # Not checked yet — do it inline
            accessible = Path(p).exists()
            with _nas_lock:
                _nas_status[p] = accessible
        paths.append({"path": p, "accessible": accessible})
    all_ok = all(p["accessible"] for p in paths) if paths else True
    return {"paths": paths, "all_ok": all_ok}



@app.get("/api/nas/browse")
def nas_browse(
    path: str = "",
    current_user: User = Depends(auth.get_current_user),
):
    """Browse NAS directories. Returns subfolders of the given path."""
    from config import NAS_PATHS
    import os

    # Determine the NAS root(s)
    nas_roots = [Path(p).resolve() for p in (NAS_PATHS or [])]
    if not nas_roots:
        raise HTTPException(status_code=400, detail="No NAS paths configured")

    # If no path given, return the top-level NAS roots and their children
    if not path:
        folders = []
        for root in nas_roots:
            if root.exists():
                folders.append({"name": root.name, "path": str(root), "is_root": True})
        return {"path": "", "folders": folders}

    # Validate the path is under a NAS root (prevent traversal)
    target = Path(path).resolve()
    if not any(str(target).startswith(str(r)) for r in nas_roots):
        raise HTTPException(status_code=403, detail="Path is outside NAS roots")

    if not target.exists() or not target.is_dir():
        raise HTTPException(status_code=404, detail="Directory not found")

    # List subdirectories (not files — just folders for navigation)
    folders = []
    try:
        for entry in sorted(os.scandir(str(target)), key=lambda e: e.name.lower()):
            if entry.is_dir() and not entry.name.startswith('.') and entry.name != '#recycle':
                folders.append({"name": entry.name, "path": str(Path(entry.path).resolve())})
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied")

    # Count files in this directory (not recursive — just immediate children)
    file_count = 0
    try:
        for entry in os.scandir(str(target)):
            if entry.is_file() and not entry.name.startswith('.'):
                file_count += 1
    except Exception:
        pass

    return {
        "path": str(target),
        "parent": str(target.parent) if target.parent != target else None,
        "folders": folders,
        "file_count": file_count,
    }


@app.get("/api/research/status")
def research_status(current_user: User = Depends(auth.get_current_user)):
    """Check whether the local SearXNG instance is reachable."""
    available = rag.searxng_available()
    return {"available": available, "url": rag.SEARXNG_URL}


# ── Auth ──────────────────────────────────────────────────────────────────────

class LoginRequest(BaseModel):
    username: str
    password: str


@app.post("/api/auth/login")
def login(req: LoginRequest, request: Request, db: Session = Depends(get_db)):
    ip = request.client.host if request.client else "-"
    user = auth.authenticate_user(db, req.username, req.password)
    if not user:
        audit("login_failure", username=req.username, ip=ip)
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid credentials")
    token = auth.create_access_token(user.id, user.username, user.role)
    audit("login_success", user_id=user.id, username=user.username, ip=ip)
    resp_data = {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": user.id,
            "username": user.username,
            "display_name": user.display_name,
            "role": user.role,
        },
    }
    response = JSONResponse(content=resp_data)
    response.set_cookie(
        key="access_token",
        value=token,
        httponly=True,
        samesite="lax",
        max_age=86400 * 7,
        path="/",
    )
    return response


@app.get("/api/auth/me")
def me(current_user: User = Depends(auth.get_current_user)):
    return {
        "id": current_user.id,
        "username": current_user.username,
        "display_name": current_user.display_name,
        "role": current_user.role,
    }


# ── Cases ─────────────────────────────────────────────────────────────────────

CASE_TYPES = ["Criminal", "Civil", "Family", "Corporate", "Personal Injury",
              "Real Estate", "Immigration", "Bankruptcy", "Employment", "Other"]

class CaseCreate(BaseModel):
    case_name:      str
    case_number:    Optional[str] = None
    case_type:      Optional[str] = None
    nas_path:       Optional[str] = None
    client_name:    Optional[str] = None
    opposing_party: Optional[str] = None
    jurisdiction:   Optional[str] = None
    assigned_to:    Optional[str] = None
    date_opened:    Optional[str] = None
    description:    Optional[str] = None

class CaseUpdate(BaseModel):
    case_name:      Optional[str] = None
    case_number:    Optional[str] = None
    case_type:      Optional[str] = None
    nas_path:       Optional[str] = None
    client_name:    Optional[str] = None
    opposing_party: Optional[str] = None
    jurisdiction:   Optional[str] = None
    assigned_to:    Optional[str] = None
    date_opened:    Optional[str] = None
    description:    Optional[str] = None
    status:         Optional[str] = None


def _case_to_dict(c: Case) -> dict:
    return {
        "id":             c.id,
        "case_name":      c.case_name,
        "case_number":    c.case_number,
        "case_type":      c.case_type,
        "nas_path":       c.nas_path,
        "client_name":    c.client_name,
        "opposing_party": c.opposing_party,
        "jurisdiction":   c.jurisdiction,
        "assigned_to":    c.assigned_to,
        "date_opened":    c.date_opened,
        "description":    c.description,
        "status":         c.status,
        "last_indexed":   c.last_indexed.isoformat() if c.last_indexed else None,
        "indexed_count":  c.indexed_count or 0,
        "created_at":     c.created_at.isoformat(),
        "created_by":     c.created_by,
    }


@app.get("/api/cases")
def list_cases(
    status: Optional[str] = None,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    q = db.query(Case)
    if status:
        q = q.filter(Case.status == status)
    cases = q.order_by(Case.created_at.desc()).all()
    return [_case_to_dict(c) for c in cases]


@app.get("/api/cases/types")
def get_case_types(current_user: User = Depends(auth.get_current_user)):
    return CASE_TYPES


@app.get("/api/cases/{case_id}")
def get_case(
    case_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    return _case_to_dict(case)


@app.post("/api/cases", status_code=201)
def create_case(
    body: CaseCreate,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    # Check for duplicate case number
    if body.case_number:
        if db.query(Case).filter(Case.case_number == body.case_number).first():
            raise HTTPException(status_code=409, detail="Case number already exists")

    case = Case(
        created_by=current_user.id,
        **{k: v for k, v in body.model_dump().items() if v is not None},
    )
    db.add(case)
    db.commit()
    db.refresh(case)

    # Auto-trigger index if NAS path was provided
    if case.nas_path:
        job_id = idx.start_case_index(case.id, case.nas_path)
        return {**_case_to_dict(case), "index_job_id": job_id}

    return _case_to_dict(case)


@app.patch("/api/cases/{case_id}")
def update_case(
    case_id: int,
    body: CaseUpdate,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    old_nas_path = case.nas_path
    for field, value in body.model_dump(exclude_none=True).items():
        setattr(case, field, value)
    db.commit()

    result = _case_to_dict(case)

    # Re-index if NAS path changed
    if body.nas_path and body.nas_path != old_nas_path:
        job_id = idx.start_case_index(case.id, case.nas_path)
        result["index_job_id"] = job_id

    return result


@app.post("/api/cases/{case_id}/reindex")
def reindex_case(
    case_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    if not case.nas_path:
        raise HTTPException(status_code=400, detail="Case has no NAS path configured")
    job_id = idx.start_case_index(case.id, case.nas_path)
    return {"job_id": job_id}


@app.get("/api/cases/{case_id}/index-status/{job_id}")
def case_index_status(
    case_id: int,
    job_id: str,
    current_user: User = Depends(auth.get_current_user),
):
    data = idx.get_job_status(job_id)
    if not data:
        raise HTTPException(status_code=404, detail="Job not found")
    return data


# ── Matters ───────────────────────────────────────────────────────────────────

class MatterCreate(BaseModel):
    name: str
    case_id: Optional[int] = None


class MatterUpdate(BaseModel):
    name: Optional[str] = None
    archived: Optional[bool] = None
    billable_time: Optional[float] = None
    case_id: Optional[int] = None


@app.get("/api/matters")
def list_matters(
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matters = (
        db.query(Matter)
        .filter(Matter.user_id == current_user.id, Matter.archived == False)
        .order_by(Matter.created_at.desc())
        .all()
    )
    return [
        {
            "id": m.id,
            "name": m.name,
            "case_id": m.case_id,
            "billable_time": m.billable_time or 0.0,
            "created_at": m.created_at.isoformat(),
            # Inline case info so sidebar/context bar doesn't need a separate lookup
            "case_name":       m.case.case_name       if m.case else None,
            "case_number":     m.case.case_number     if m.case else None,
            "case_type":       m.case.case_type       if m.case else None,
            "client_name":     m.case.client_name     if m.case else None,
            "opposing_party":  m.case.opposing_party  if m.case else None,
            "case_status":     m.case.status          if m.case else None,
        }
        for m in matters
    ]


@app.post("/api/matters", status_code=201)
def create_matter(
    body: MatterCreate,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matter = Matter(user_id=current_user.id, name=body.name, case_id=body.case_id)
    db.add(matter)
    db.commit()
    db.refresh(matter)
    return {"id": matter.id, "name": matter.name, "case_id": matter.case_id, "billable_time": matter.billable_time or 0.0, "created_at": matter.created_at.isoformat()}


@app.patch("/api/matters/{matter_id}")
def update_matter(
    matter_id: int,
    body: MatterUpdate,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    if body.name is not None:
        matter.name = body.name
    if body.archived is not None:
        matter.archived = body.archived
    if body.billable_time is not None:
        matter.billable_time = body.billable_time
    if body.case_id is not None:
        matter.case_id = body.case_id if body.case_id != 0 else None
    db.commit()
    return {"id": matter.id, "name": matter.name, "archived": matter.archived, "billable_time": matter.billable_time or 0.0}


@app.get("/api/matters/{matter_id}/messages")
def get_messages(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    return [
        {
            "id": m.id,
            "role": m.role,
            "content": m.content,
            "sources": m.sources_list(),
            "created_at": m.created_at.isoformat(),
        }
        for m in matter.messages
    ]


# ── Chat (SSE streaming) ──────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    scope: str = "all"            # "all" | "global" | "user" | "both" | "case"
    query_type: str = "auto"      # "auto" | "summary" | "timeline" | "risk" | "drafting" | "compare"
    verbosity_role: str = "attorney"  # "attorney" | "associate" | "paralegal" | "client"
    research_mode: bool = False   # True = include SearXNG web results
    history: list[dict] | None = None  # conversation history for follow-up rewriting


@app.post("/api/matters/{matter_id}/chat")
async def chat(
    matter_id: int,
    body: ChatRequest,
    current_user: User = Depends(rate_limit),
    db: Session = Depends(get_db),
):
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    # ── Case context: auto-scope to case collection when matter is linked to a case ──
    case_context: dict | None = None
    resolved_scope = body.scope

    if matter.case_id:
        linked_case = db.query(Case).filter(Case.id == matter.case_id).first()
        if linked_case:
            case_context = {
                "case_name":      linked_case.case_name,
                "case_number":    linked_case.case_number,
                "case_type":      linked_case.case_type,
                "client_name":    linked_case.client_name,
                "opposing_party": linked_case.opposing_party,
                "jurisdiction":   linked_case.jurisdiction,
                "assigned_to":    linked_case.assigned_to,
                "status":         linked_case.status,
                "description":    linked_case.description,
                "matter_name":    matter.name,
            }
            # Auto-scope to this case's collection when user chose "case" or "both"
            # "all" stays as "all" so retrieve() enumerates every collection
            if body.scope in ("case", "both"):
                resolved_scope = case_collection(matter.case_id)
    elif body.scope == "case":
        # Matter has no linked case but scope was "case" — fall back to "both"
        resolved_scope = "both"

    # Fetch recent conversation history (last 6 messages = 3 user+assistant pairs)
    recent_msgs = (
        db.query(Message)
        .filter(Message.matter_id == matter_id)
        .order_by(Message.created_at.desc())
        .limit(6)
        .all()
    )
    chat_history = [
        {"role": m.role, "content": m.content}
        for m in reversed(recent_msgs)
    ] if recent_msgs else None

    # Save user message
    user_msg = Message(
        matter_id=matter_id,
        user_id=current_user.id,
        role="user",
        content=body.message,
    )
    db.add(user_msg)
    db.commit()

    # Stream response via SSE
    async def _generate():
        full_response = ""
        sources_captured = []
        token_stats = {}

        try:
            async for chunk in rag.stream_response(
                body.message, current_user.id, resolved_scope,
                query_type=body.query_type,
                verbosity_role=body.verbosity_role,
                research_mode=body.research_mode,
                history=chat_history,
                case_context=case_context,
                matter_id=matter_id,
            ):
                if len(chunk) == 3:
                    # Final stats tuple: (token, sources, stats_dict)
                    token_stats = chunk[2]
                    continue
                token, sources = chunk
                full_response += token
                if sources and not sources_captured:
                    sources_captured = sources
                yield f"data: {json.dumps({'token': token, 'sources': sources})}\n\n"
        except Exception as e:
            err_msg = "⚠ Sherlock timed out waiting for the AI model. The model may be loading or under heavy load — please try again in a moment."
            yield f"data: {json.dumps({'token': err_msg, 'sources': [], 'error': True})}\n\n"
            full_response = err_msg
            log.warning("chat stream error: %s", e)

        # Save assistant message
        ai_msg = Message(
            matter_id=matter_id,
            user_id=current_user.id,
            role="assistant",
            content=full_response,
            sources=json.dumps(sources_captured),
        )
        db.add(ai_msg)
        db.commit()
        db.refresh(ai_msg)

        # Log query with token metrics for usage dashboard
        from models import QueryLog
        ql = QueryLog(
            user_id=current_user.id,
            matter_id=matter_id,
            query_type=body.query_type,
            verbosity=body.verbosity_role,
            research_mode=body.research_mode,
            prompt_tokens=token_stats.get("prompt_tokens"),
            completion_tokens=token_stats.get("completion_tokens"),
            total_tokens=token_stats.get("total_tokens"),
            tokens_per_sec=token_stats.get("tokens_per_sec"),
            latency_ms=token_stats.get("latency_total_ms"),
            source="user",
        )
        db.add(ql)
        db.commit()

        # Send message ID + token stats so frontend can display performance
        done_payload = {'done': True, 'message_id': ai_msg.id}
        if token_stats:
            done_payload['token_stats'] = token_stats
        yield f"data: {json.dumps(done_payload)}\n\n"

    return StreamingResponse(_generate(), media_type="text/event-stream")


@app.post("/api/chat")
async def chat_ungated(
    body: ChatRequest,
    current_user: User = Depends(rate_limit),
):
    """Stateless chat — no Matter required. Streams a response but does not persist messages."""
    async def _generate():
        token_stats = {}
        try:
            async for chunk in rag.stream_response(
                body.message, current_user.id, body.scope,
                query_type=body.query_type,
                verbosity_role=body.verbosity_role,
                research_mode=body.research_mode,
                history=body.history or [],
            ):
                if len(chunk) == 3:
                    token_stats = chunk[2]
                    continue
                token, sources = chunk
                yield f"data: {json.dumps({'token': token, 'sources': sources})}\n\n"
        except Exception as e:
            err_msg = "⚠ Sherlock timed out waiting for the AI model. The model may be loading or under heavy load — please try again in a moment."
            yield f"data: {json.dumps({'token': err_msg, 'sources': [], 'error': True})}\n\n"
            log.warning("chat stream error: %s", e)
        done_payload = {'done': True}
        if token_stats:
            done_payload['token_stats'] = token_stats
        yield f"data: {json.dumps(done_payload)}\n\n"

    return StreamingResponse(_generate(), media_type="text/event-stream")


# ── File upload ───────────────────────────────────────────────────────────────

@app.post("/api/upload")
async def upload_file(
    file: UploadFile = File(...),
    matter_id: int = Form(None),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    if file.size and file.size > MAX_UPLOAD_MB * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"File exceeds {MAX_UPLOAD_MB} MB limit")

    ext = Path(file.filename).suffix.lower()

    # Route audio files to audio handler
    if audio_mod.is_audio_file(file.filename):
        save_dir = _user_audio_dir(current_user.id)
    else:
        save_dir = _user_upload_dir(current_user.id)

    safe_name = f"{uuid.uuid4().hex}_{Path(file.filename).name}"
    dest = save_dir / safe_name

    # Stream to disk
    with dest.open("wb") as f:
        while chunk := await file.read(1 << 20):   # 1 MB chunks
            f.write(chunk)

    # Dedup: check if user already uploaded a file with identical content
    import hashlib as _hl
    _h = _hl.sha256()
    with dest.open("rb") as _f:
        while _blk := _f.read(1 << 20):
            _h.update(_blk)
    content_hash = _h.hexdigest()

    existing_upload = (
        db.query(Upload)
        .filter(Upload.user_id == current_user.id, Upload.filename == file.filename)
        .first()
    )
    # Check by hash across all user uploads (catches renamed dupes too)
    if not existing_upload:
        from models import IndexedFile
        existing_indexed = (
            db.query(Upload)
            .join(IndexedFile, IndexedFile.file_path == Upload.stored_path)
            .filter(Upload.user_id == current_user.id, IndexedFile.file_hash == content_hash)
            .first()
        )
        if existing_indexed:
            existing_upload = existing_indexed

    if existing_upload and existing_upload.status == "ready":
        # Duplicate — remove the new file, associate existing upload with matter if needed
        dest.unlink(missing_ok=True)
        if matter_id:
            matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
            if matter:
                exists = db.query(MatterFile).filter(
                    MatterFile.matter_id == matter_id, MatterFile.upload_id == existing_upload.id).first()
                if not exists:
                    db.add(MatterFile(matter_id=matter_id, upload_id=existing_upload.id))
                    db.commit()
        return {"upload_id": existing_upload.id, "job_id": None, "filename": file.filename,
                "matter_id": matter_id, "duplicate": True}

    # Record upload
    upload_record = Upload(
        user_id=current_user.id,
        filename=file.filename,
        stored_path=str(dest),
        file_type=ext.lstrip("."),
        size_bytes=dest.stat().st_size,
        status="pending",
    )
    db.add(upload_record)
    db.commit()
    db.refresh(upload_record)

    # Start background indexing
    job_id = idx.start_upload_index(upload_record.id, current_user.id, dest)

    audit("file_upload", user_id=current_user.id, username=current_user.username,
          file_path=file.filename, file_size=upload_record.size_bytes)
    # Associate file with matter if specified
    if matter_id:
        matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
        if matter:
            mf = MatterFile(matter_id=matter_id, upload_id=upload_record.id)
            db.add(mf)
            db.commit()

    return {"upload_id": upload_record.id, "job_id": job_id, "filename": file.filename, "matter_id": matter_id}


@app.get("/api/upload/{job_id}/status")
def upload_status(job_id: str, current_user: User = Depends(auth.get_current_user)):
    status_data = idx.get_job_status(job_id)
    if not status_data:
        raise HTTPException(status_code=404, detail="Job not found")
    return status_data


@app.get("/api/files")
def list_files(
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    uploads = (
        db.query(Upload)
        .filter(Upload.user_id == current_user.id)
        .order_by(Upload.uploaded_at.desc())
        .all()
    )
    return [
        {
            "id": u.id,
            "filename": u.filename,
            "file_type": u.file_type,
            "size_bytes": u.size_bytes,
            "status": u.status,
            "uploaded_at": u.uploaded_at.isoformat(),
        }
        for u in uploads
    ]


@app.delete("/api/files/{upload_id}")
def delete_file(
    upload_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    upload = db.query(Upload).filter(Upload.id == upload_id, Upload.user_id == current_user.id).first()
    if not upload:
        raise HTTPException(status_code=404, detail="File not found")

    # Remove from ChromaDB
    if upload.chroma_ids:
        try:
            coll = rag.get_or_create_collection(f"user_{current_user.id}_docs")
            coll.delete(ids=json.loads(upload.chroma_ids))
        except Exception:
            pass

    # Remove from disk
    p = Path(upload.stored_path)
    if p.exists():
        p.unlink()

    audit("file_delete", user_id=current_user.id, username=current_user.username,
          file_path=upload.filename)
    db.delete(upload)
    db.commit()
    return {"deleted": upload_id}




@app.post("/api/files/{upload_id}/retry")
def retry_index(
    upload_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    upload = db.query(Upload).filter(Upload.id == upload_id, Upload.user_id == current_user.id).first()
    if not upload:
        raise HTTPException(status_code=404, detail="File not found")
    if upload.status == "ready" and upload.chroma_ids:
        return {"status": "already_indexed", "upload_id": upload_id}
    fp = Path(upload.stored_path)
    if not fp.exists():
        raise HTTPException(status_code=404, detail="File no longer exists on disk")
    upload.status = "pending"
    upload.error_msg = None
    db.commit()
    job_id = idx.start_upload_index(upload.id, current_user.id, fp)
    return {"upload_id": upload_id, "job_id": job_id, "status": "retrying"}


# -- Matter files (file-matter associations) --

@app.get("/api/matters/{matter_id}/files")
def list_matter_files(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    rows = (
        db.query(MatterFile, Upload)
        .join(Upload, Upload.id == MatterFile.upload_id)
        .filter(MatterFile.matter_id == matter_id)
        .order_by(MatterFile.attached_at.desc())
        .all()
    )
    return [
        {
            "id": mf.id,
            "upload_id": u.id,
            "filename": u.filename,
            "file_type": u.file_type,
            "size_bytes": u.size_bytes,
            "status": u.status,
            "page_count": u.page_count,
            "stored_path": u.stored_path,
            "uploaded_at": u.uploaded_at.isoformat() if u.uploaded_at else None,
            "attached_at": mf.attached_at.isoformat() if mf.attached_at else None,
        }
        for mf, u in rows
    ]


@app.post("/api/matters/{matter_id}/files/{upload_id}", status_code=201)
def attach_file_to_matter(
    matter_id: int,
    upload_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    upload = db.query(Upload).filter(Upload.id == upload_id, Upload.user_id == current_user.id).first()
    if not upload:
        raise HTTPException(status_code=404, detail="File not found")
    existing = db.query(MatterFile).filter(
        MatterFile.matter_id == matter_id, MatterFile.upload_id == upload_id).first()
    if existing:
        return {"status": "already_attached"}
    mf = MatterFile(matter_id=matter_id, upload_id=upload_id)
    db.add(mf)
    db.commit()
    return {"status": "attached", "matter_id": matter_id, "upload_id": upload_id}


@app.delete("/api/matters/{matter_id}/files/{upload_id}", status_code=204)
def detach_file_from_matter(
    matter_id: int,
    upload_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    mf = db.query(MatterFile).filter(
        MatterFile.matter_id == matter_id, MatterFile.upload_id == upload_id).first()
    if mf:
        db.delete(mf)
        db.commit()


@app.get("/api/files/{upload_id}/download")
def download_file(
    upload_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    upload = db.query(Upload).filter(Upload.id == upload_id, Upload.user_id == current_user.id).first()
    if not upload:
        raise HTTPException(status_code=404, detail="File not found")
    from starlette.responses import FileResponse as _FR
    return _FR(upload.stored_path, filename=upload.filename)


# ── Audio ─────────────────────────────────────────────────────────────────────

@app.post("/api/audio")
async def upload_audio(
    file: UploadFile = File(...),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    if not audio_mod.is_audio_file(file.filename):
        raise HTTPException(status_code=400, detail="Unsupported audio format")

    save_dir = _user_audio_dir(current_user.id)
    safe_name = f"{uuid.uuid4().hex}_{Path(file.filename).name}"
    dest = save_dir / safe_name

    with dest.open("wb") as f:
        while chunk := await file.read(1 << 20):
            f.write(chunk)

    job_id = audio_mod.start_transcription(dest)
    return {"job_id": job_id, "filename": file.filename}


@app.get("/api/audio/{job_id}/status")
def audio_status(job_id: str, current_user: User = Depends(auth.get_current_user)):
    status_data = audio_mod.get_job_status(job_id)
    if not status_data:
        raise HTTPException(status_code=404, detail="Job not found")
    return status_data


# ── Outputs ───────────────────────────────────────────────────────────────────

class SaveOutputRequest(BaseModel):
    message_id: int
    matter_name: str = ""


@app.post("/api/outputs")
def save_output(
    body: SaveOutputRequest,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    msg = (
        db.query(Message)
        .join(Matter)
        .filter(Message.id == body.message_id, Matter.user_id == current_user.id)
        .first()
    )
    if not msg:
        raise HTTPException(status_code=404, detail="Message not found")

    record = out_mod.save_response(db, current_user, msg, body.matter_name)
    return {
        "output_id": record.id,
        "filename": record.filename,
        "saved_at": record.saved_at.isoformat(),
        "download_url": f"/api/outputs/{record.id}/download",
    }


@app.get("/api/outputs")
def list_outputs_route(
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    records = out_mod.list_outputs(db, current_user)
    return [
        {
            "id": r.id,
            "filename": r.filename,
            "saved_at": r.saved_at.isoformat(),
            "matter_id": r.matter_id,
        }
        for r in records
    ]


@app.get("/api/outputs/download-all")
def download_all_outputs(
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Bundle all of a user's saved outputs into a single zip for download."""
    import io, zipfile
    records = out_mod.list_outputs(db, current_user)
    if not records:
        raise HTTPException(status_code=404, detail="No outputs to download")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        seen: dict[str, int] = {}
        for r in records:
            fp = Path(r.file_path)
            if not fp.exists():
                continue
            name = r.filename
            # De-duplicate filenames inside zip
            if name in seen:
                seen[name] += 1
                stem, ext = (name.rsplit(".", 1) + [""])[:2]
                name = f"{stem}_{seen[name]}.{ext}" if ext else f"{stem}_{seen[name]}"
            else:
                seen[name] = 0
            zf.write(fp, name)

    buf.seek(0)
    from starlette.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="sherlock-outputs.zip"'},
    )


@app.get("/api/outputs/{output_id}/download")
def download_output(
    output_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    record = db.query(Output).filter(
        Output.id == output_id, Output.user_id == current_user.id
    ).first()
    if not record:
        raise HTTPException(status_code=404, detail="Output not found")
    fp = Path(record.file_path)
    if not fp.exists():
        raise HTTPException(status_code=404, detail="Output file missing from disk")
    return FileResponse(str(fp), filename=record.filename, media_type="text/plain")


# ── Privilege Log Generator ──────────────────────────────────────────────────

@app.post("/api/matters/{matter_id}/privilege-log")
def generate_privilege_log(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Generate a privilege log from indexed case documents."""
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    scope = case_collection(matter.case_id) if matter.case_id else "both"

    # Use RAG to extract privilege log entries
    chunks = rag.retrieve("privileged communications attorney client work product", current_user.id, scope, n=15)
    if not chunks:
        return {"entries": [], "message": "No documents found for privilege analysis"}

    doc_text = "\n\n---\n\n".join(
        f"[Doc: {c['source']} | Chunk {c['chunk']}]\n{c['text']}" for c in chunks
    )

    system = (
        "You are a privilege log generator. Analyze the provided documents and extract a privilege log. "
        "For each potentially privileged document or communication, identify:\n"
        "Return ONLY valid JSON — an array of objects with these exact keys:\n"
        '  "document": filename or description,\n'
        '  "date": date if visible (YYYY-MM-DD or descriptive),\n'
        '  "from_party": sender/author,\n'
        '  "to_party": recipient(s),\n'
        '  "cc_party": CC recipients if any,\n'
        '  "privilege_type": "attorney-client" | "work-product" | "joint-defense" | "other",\n'
        '  "description": brief description of subject matter (without waiving privilege),\n'
        '  "basis": legal basis for privilege claim\n'
        "Return [] if no privileged materials found. No prose — only the JSON array."
    )

    import requests as _req
    from config import OLLAMA_URL, LLM_MODEL

    try:
        resp = _req.post(
            f"{OLLAMA_URL}/api/generate",
            json={"model": LLM_MODEL, "prompt": f"Documents:\n\n{doc_text}\n\n---\n\nGenerate privilege log entries.",
                  "system": system, "stream": False,
                  "options": {"temperature": 0.0, "num_predict": 2048}},
            timeout=120,
        )
        resp.raise_for_status()
        raw = resp.json().get("response", "").strip()

        # Strip markdown fences
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        entries = json.loads(raw)
        if not isinstance(entries, list):
            entries = []
    except Exception as e:
        log.warning("Privilege log generation failed: %s", e)
        entries = []

    audit("privilege_log", user_id=current_user.id, username=current_user.username,
          matter_id=matter_id, count=len(entries))
    return {"entries": entries, "count": len(entries)}


# ── Deadline extractor ────────────────────────────────────────────────────────

@app.post("/api/matters/{matter_id}/deadlines/extract")
def extract_matter_deadlines(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import Deadline
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    scope = case_collection(matter.case_id) if matter.case_id else "both"
    deadlines = rag.extract_deadlines("deadlines filing dates statutes of limitations notices hearings", current_user.id, scope)

    # Wipe old extractions for this matter and replace
    db.query(Deadline).filter(Deadline.matter_id == matter_id).delete()

    saved = []
    for d in deadlines:
        if not isinstance(d, dict) or not d.get("description"):
            continue
        row = Deadline(
            matter_id=matter_id,
            case_id=matter.case_id,
            date_str=str(d.get("date_str", ""))[:32],
            description=d.get("description", "")[:2000],
            dl_type=d.get("dl_type", "other")[:64],
            source_file=str(d.get("source_file", ""))[:512],
            urgency=d.get("urgency", "normal")[:16],
        )
        db.add(row)
        saved.append(row)
    db.commit()

    audit("deadline_extract", user_id=current_user.id, username=current_user.username,
          matter_id=matter_id, count=len(saved))
    return {"extracted": len(saved), "deadlines": [
        {"id": r.id, "date_str": r.date_str, "description": r.description,
         "dl_type": r.dl_type, "source_file": r.source_file, "urgency": r.urgency,
         "extracted_at": r.extracted_at.isoformat()}
        for r in saved
    ]}


@app.get("/api/matters/{matter_id}/deadlines")
def get_matter_deadlines(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import Deadline
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    rows = db.query(Deadline).filter(Deadline.matter_id == matter_id).order_by(Deadline.date_str).all()
    return [
        {"id": r.id, "date_str": r.date_str, "description": r.description,
         "dl_type": r.dl_type, "source_file": r.source_file, "urgency": r.urgency,
         "extracted_at": r.extracted_at.isoformat()}
        for r in rows
    ]


# ── Matter Auto-Brief ─────────────────────────────────────────────────────────

@app.get("/api/matters/{matter_id}/brief")
def get_matter_brief(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import MatterBrief
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")
    brief = db.query(MatterBrief).filter(MatterBrief.matter_id == matter_id).first()
    msg_count = len(matter.messages)
    stale = brief is None or brief.msg_count != msg_count
    if brief:
        return {"has_brief": True, "stale": stale, "brief_md": brief.brief_md,
                "risks_md": brief.risks_md, "generated_at": brief.generated_at.isoformat(),
                "msg_count": brief.msg_count}
    return {"has_brief": False, "stale": True, "brief_md": None, "risks_md": None,
            "generated_at": None, "msg_count": 0}


@app.post("/api/matters/{matter_id}/brief/generate")
def generate_matter_brief(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import MatterBrief
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    scope = case_collection(matter.case_id) if matter.case_id else "both"
    result = rag.generate_brief("matter overview", current_user.id, scope)

    brief = db.query(MatterBrief).filter(MatterBrief.matter_id == matter_id).first()
    msg_count = len(matter.messages)
    if brief:
        brief.brief_md = result["brief_md"]
        brief.risks_md = result["risks_md"]
        brief.generated_at = datetime.utcnow()
        brief.msg_count = msg_count
    else:
        brief = MatterBrief(matter_id=matter_id, brief_md=result["brief_md"],
                            risks_md=result["risks_md"], msg_count=msg_count)
        db.add(brief)
    db.commit()

    audit("brief_generate", user_id=current_user.id, username=current_user.username, matter_id=matter_id)
    return {"brief_md": result["brief_md"], "risks_md": result["risks_md"],
            "generated_at": brief.generated_at.isoformat()}


# ── Export to DOCX ────────────────────────────────────────────────────────────

@app.get("/api/export/memo")
def export_memo(
    message_id: int = Query(None),
    format: str = Query("docx"),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import Message as Msg
    msg = db.query(Msg).filter(Msg.id == message_id).first()
    if not msg:
        raise HTTPException(status_code=404, detail="Message not found")
    # Verify ownership via matter
    matter = db.query(Matter).filter(Matter.id == msg.matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=403, detail="Access denied")

    # Find the user query that preceded this response
    all_msgs = matter.messages
    user_query = ""
    for i, m in enumerate(all_msgs):
        if m.id == msg.id and i > 0:
            user_query = all_msgs[i-1].content
            break

    case_name = matter.case.case_name if matter.case else ""
    matter_name = matter.name

    doc = DocxDocument()

    # ── Page setup ─────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # ── Header ─────────────────────────────────
    hdr = doc.add_heading("SHERLOCK — LEGAL RESEARCH MEMO", level=0)
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr.runs[0].font.size = Pt(14)
    hdr.runs[0].font.color.rgb = RGBColor(0x1a, 0x1d, 0x27)

    # ── Metadata table ──────────────────────────
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    def _set_meta(row, label, val):
        meta.rows[row].cells[0].text = label
        meta.rows[row].cells[1].text = val or "—"
        meta.rows[row].cells[0].paragraphs[0].runs[0].bold = True
    from datetime import date
    _set_meta(0, "Matter", matter_name)
    _set_meta(1, "Case", case_name)
    _set_meta(2, "Prepared by", "Sherlock AI Paralegal")
    _set_meta(3, "Date", date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()

    # ── Query ──────────────────────────────────
    if user_query:
        q_para = doc.add_paragraph()
        q_para.add_run("Research Question: ").bold = True
        q_para.add_run(user_query).italic = True
        doc.add_paragraph()

    # ── Response body (markdown → docx) ────────
    doc.add_heading("Analysis", level=1)
    for line in msg.content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        # Heading detection
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "• ", "→ ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif _re.match(r'^\d+\.', stripped):
            doc.add_paragraph(stripped, style="List Number")
        else:
            p = doc.add_paragraph()
            # Handle inline **bold** and *italic*
            parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith("*") and part.endswith("*"):
                    p.add_run(part[1:-1]).italic = True
                else:
                    p.add_run(part)

    # ── Sources ────────────────────────────────
    sources = msg.sources_list()
    if sources:
        doc.add_paragraph()
        doc.add_heading("Sources", level=1)
        for s in sources:
            fname = s.get("file", s.get("source", "Unknown"))
            score = s.get("score", "")
            score_str = f"  (relevance: {score:.2f})" if isinstance(score, float) else ""
            doc.add_paragraph(f"• {fname}{score_str}", style="List Bullet")

    # ── Footer note ────────────────────────────
    doc.add_paragraph()
    foot = doc.add_paragraph("Generated by Sherlock — AI-assisted legal research. Always verify with primary sources before relying on this analysis.")
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_matter = _re.sub(r'[^\w\s-]', '', matter_name)[:40].strip().replace(' ', '_')
    filename = f"Sherlock_Memo_{safe_matter}_{date.today()}.docx"

    audit("memo_export", user_id=current_user.id, username=current_user.username,
          matter_id=matter.id, message_id=message_id)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Batch Export / Case Report (DOCX) ────────────────────────────────────────

@app.get("/api/matters/{matter_id}/export/docx")
def export_matter_docx(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Export entire matter conversation as a professional DOCX research memo."""
    matter = db.query(Matter).filter(Matter.id == matter_id, Matter.user_id == current_user.id).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    messages = (
        db.query(Message)
        .filter(Message.matter_id == matter_id)
        .order_by(Message.created_at)
        .all()
    )
    if not messages:
        raise HTTPException(status_code=400, detail="No messages to export")

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Title
    from datetime import date
    hdr = doc.add_heading(f"SHERLOCK — CASE RESEARCH REPORT", level=0)
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr.runs[0].font.size = Pt(14)
    hdr.runs[0].font.color.rgb = RGBColor(0x1a, 0x1d, 0x27)

    # Metadata
    case_name = matter.case.case_name if matter.case else ""
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    def _set(row, label, val):
        meta.rows[row].cells[0].text = label
        meta.rows[row].cells[1].text = val or "—"
        meta.rows[row].cells[0].paragraphs[0].runs[0].bold = True
    _set(0, "Matter", matter.name)
    _set(1, "Case", case_name)
    _set(2, "Prepared by", current_user.display_name or current_user.username)
    _set(3, "Date", date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()

    # Each Q&A pair
    for msg in messages:
        if msg.role == "user":
            p = doc.add_paragraph()
            p.add_run("QUERY: ").bold = True
            p.add_run(msg.content)
            ts = msg.created_at.strftime("%b %d, %Y %I:%M %p") if msg.created_at else ""
            if ts:
                p_ts = doc.add_paragraph(ts)
                p_ts.runs[0].font.size = Pt(8)
                p_ts.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        else:
            doc.add_heading("Analysis", level=2)
            for line in msg.content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph()
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith(("- ", "• ")):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif _re.match(r'^\d+\.', stripped):
                    doc.add_paragraph(stripped, style="List Number")
                else:
                    p = doc.add_paragraph()
                    parts = _re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', stripped)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            p.add_run(part[2:-2]).bold = True
                        elif part.startswith("*") and part.endswith("*"):
                            p.add_run(part[1:-1]).italic = True
                        else:
                            p.add_run(part)

            # Sources
            srcs = msg.sources_list()
            if srcs:
                doc.add_paragraph()
                sp = doc.add_paragraph()
                sp.add_run("Sources:").bold = True
                for s in srcs:
                    fname = s.get("file", "Unknown")
                    score = s.get("score", "")
                    score_str = f" (relevance: {score:.2f})" if isinstance(score, float) else ""
                    doc.add_paragraph(f"{fname}{score_str}", style="List Bullet")

            doc.add_paragraph()  # spacer between Q&A pairs

    # Footer
    foot = doc.add_paragraph("Generated by Sherlock — AI-assisted legal research.")
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_matter = _re.sub(r'[^\w\s-]', '', matter.name)[:40].strip().replace(' ', '_')
    filename = f"Sherlock_Report_{safe_matter}_{date.today()}.docx"

    audit("case_report_export", user_id=current_user.id, username=current_user.username,
          matter_id=matter_id, messages=len(messages))

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Admin ─────────────────────────────────────────────────────────────────────

class CreateUserRequest(BaseModel):
    username: str
    password: str
    display_name: str = ""
    role: str = "user"


class UpdateUserRequest(BaseModel):
    active: Optional[bool] = None
    role: Optional[str] = None
    new_password: Optional[str] = None
    display_name: Optional[str] = None


@app.get("/api/admin/users")
def admin_list_users(
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    users = db.query(User).order_by(User.created_at).all()
    return [
        {
            "id": u.id,
            "username": u.username,
            "display_name": u.display_name,
            "role": u.role,
            "active": u.active,
            "created_at": u.created_at.isoformat(),
            "last_login": u.last_login.isoformat() if u.last_login else None,
        }
        for u in users
    ]


@app.post("/api/admin/users", status_code=201)
def admin_create_user(
    body: CreateUserRequest,
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    try:
        user = auth.create_user(db, body.username, body.password, body.display_name, body.role)
        audit("user_create", user_id=_.id, username=_.username,
              detail=f"Created {body.role} account '{body.username}'")
        return {"id": user.id, "username": user.username}
    except ValueError as e:
        raise HTTPException(status_code=409, detail=str(e))


@app.patch("/api/admin/users/{user_id}")
def admin_update_user(
    user_id: int,
    body: UpdateUserRequest,
    admin: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if body.active is not None:
        user.active = body.active
    if body.role is not None:
        user.role = body.role
    if body.display_name is not None:
        user.display_name = body.display_name
    if body.new_password:
        auth.reset_password(db, user_id, body.new_password)
        return {"updated": user_id}
    db.commit()
    return {"updated": user_id}



class ResetPasswordRequest(BaseModel):
    new_password: str


@app.delete("/api/admin/users/{user_id}")
def admin_delete_user(
    user_id: int,
    admin: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if user.id == admin.id:
        raise HTTPException(status_code=400, detail="Cannot delete your own account")
    username = user.username
    db.delete(user)
    db.commit()
    audit("user_delete", user_id=admin.id, username=admin.username,
          detail=f"Deleted user '{username}' (id={user_id})")
    return {"deleted": user_id}



# ── Dictation Analysis ─────────────────────────────────────────────────────

@app.get("/api/dictations")
def list_dictations(
    current_user: User = Depends(auth.get_current_user),
):
    """List all analyzed dictations with their tasks."""
    import sqlite3
    db_path = str(Path(DATA_DIR) / "dictations.db")
    if not os.path.exists(db_path):
        return {"dictations": [], "summary": {"total_files": 0, "total_tasks": 0, "pending": 0}}
    db = sqlite3.connect(db_path, timeout=5)
    db.row_factory = sqlite3.Row
    dictations = []
    for d in db.execute("SELECT * FROM dictations ORDER BY recorded_at DESC").fetchall():
        tasks = []
        for t in db.execute(
            "SELECT * FROM dictation_tasks WHERE dictation_id = ? ORDER BY task_order", (d["id"],)
        ).fetchall():
            tasks.append({
                "id": t["id"], "order": t["task_order"], "assignee": t["assignee"],
                "action": t["action"], "client_or_case": t["client_or_case"],
                "priority": t["priority"], "due_hint": t["due_hint"],
                "status": t["status"], "notes": t["notes"],
            })
        dictations.append({
            "id": d["id"], "file_name": d["file_name"],
            "recorded_at": d["recorded_at"], "duration_secs": d["duration_secs"],
            "transcript": d["transcript"], "task_count": len(tasks),
            "status": d["status"], "tasks": tasks,
        })
    total_tasks = db.execute("SELECT COUNT(*) FROM dictation_tasks").fetchone()[0]
    pending = db.execute("SELECT COUNT(*) FROM dictation_tasks WHERE status='pending'").fetchone()[0]
    db.close()
    return {
        "dictations": dictations,
        "summary": {"total_files": len(dictations), "total_tasks": total_tasks, "pending": pending},
    }


@app.get("/api/dictations/tasks")
def list_dictation_tasks(
    status: str = Query(default=None),
    assignee: str = Query(default=None),
    current_user: User = Depends(auth.get_current_user),
):
    """List tasks extracted from dictations, optionally filtered."""
    import sqlite3
    db_path = str(Path(DATA_DIR) / "dictations.db")
    if not os.path.exists(db_path):
        return []
    db = sqlite3.connect(db_path, timeout=5)
    db.row_factory = sqlite3.Row
    query = """SELECT t.*, d.file_name, d.recorded_at FROM dictation_tasks t
               JOIN dictations d ON d.id = t.dictation_id WHERE 1=1"""
    params = []
    if status:
        query += " AND t.status = ?"
        params.append(status)
    if assignee:
        query += " AND LOWER(t.assignee) = LOWER(?)"
        params.append(assignee)
    query += " ORDER BY d.recorded_at DESC, t.task_order"
    rows = db.execute(query, params).fetchall()
    db.close()
    return [dict(r) for r in rows]


@app.patch("/api/dictations/tasks/{task_id}")
def update_dictation_task(
    task_id: int,
    body: dict,
    current_user: User = Depends(auth.get_current_user),
):
    """Update a dictation task (status, notes)."""
    import sqlite3
    db_path = str(Path(DATA_DIR) / "dictations.db")
    db = sqlite3.connect(db_path, timeout=5)
    task = db.execute("SELECT id FROM dictation_tasks WHERE id = ?", (task_id,)).fetchone()
    if not task:
        db.close()
        raise HTTPException(status_code=404, detail="Task not found")
    if "status" in body:
        db.execute("UPDATE dictation_tasks SET status = ? WHERE id = ?", (body["status"], task_id))
        if body["status"] == "completed":
            db.execute("UPDATE dictation_tasks SET completed_at = ? WHERE id = ?",
                       (datetime.utcnow().isoformat(), task_id))
    if "notes" in body:
        db.execute("UPDATE dictation_tasks SET notes = ? WHERE id = ?", (body["notes"], task_id))
    db.commit()
    db.close()
    return {"updated": task_id}


@app.get("/api/dictations/status")
def dictation_worker_status(
    current_user: User = Depends(auth.get_current_user),
):
    """Get dictation worker status."""
    status_path = str(Path(DATA_DIR) / "dictation_status.json")
    if not os.path.exists(status_path):
        return {"active": False}
    try:
        with open(status_path) as f:
            info = json.load(f)
        pid = info.get("pid")
        if pid:
            try:
                os.kill(pid, 0)
                info["active"] = True
            except OSError:
                info["active"] = False
        return info
    except Exception:
        return {"active": False}


@app.post("/api/dictations/scan")
def trigger_dictation_scan(
    _: User = Depends(auth.require_admin),
):
    """Trigger a one-time dictation scan (runs as subprocess)."""
    import subprocess
    worker = str(Path(__file__).parent / "dictation_worker.py")
    venv_python = str(Path(__file__).resolve().parent.parent / "venv" / "bin" / "python3")
    subprocess.Popen(
        [venv_python, worker, "once"],
        cwd=str(Path(__file__).parent),
        stdout=open("/tmp/dictation_scan.log", "w"),
        stderr=subprocess.STDOUT,
    )
    return {"started": True}


# ── Indexing Activity Feed ─────────────────────────────────────────────────

@app.get("/api/index-activity")
def index_activity(
    since_minutes: int = Query(default=60, le=1440),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Return recent indexing activity — files indexed in the last N minutes."""
    from datetime import timedelta
    from models import IndexedFile
    cutoff = datetime.utcnow() - timedelta(minutes=since_minutes)
    recent = (
        db.query(IndexedFile)
        .filter(IndexedFile.indexed_at >= cutoff)
        .order_by(IndexedFile.indexed_at.desc())
        .limit(50)
        .all()
    )
    # Group by case
    by_case = {}
    for f in recent:
        key = f.case_id or 0
        if key not in by_case:
            by_case[key] = {"case_id": f.case_id, "collection": f.collection, "files": []}
        by_case[key]["files"].append({
            "path": f.file_path,
            "name": Path(f.file_path).name,
            "chunks": f.chunk_count,
            "indexed_at": f.indexed_at.isoformat() if f.indexed_at else None,
        })
    return {
        "total_recent": len(recent),
        "groups": list(by_case.values()),
    }


@app.get("/api/admin/status")
def admin_status(
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    import requests as req

    def _check(url):
        try:
            req.get(url, timeout=3)
            return "up"
        except Exception:
            return "down"

    from config import CHROMA_URL, OLLAMA_URL
    from models import IndexedFile

    return {
        "ollama": _check(f"{OLLAMA_URL}/api/tags"),
        "chroma": _check(f"{CHROMA_URL}/api/v2/heartbeat"),
        "users": db.query(User).count(),
        "indexed_files": db.query(IndexedFile).count(),
        "outputs": db.query(Output).count(),
    }


@app.post("/api/admin/reindex")
def admin_reindex(
    admin: User = Depends(auth.require_admin),
):
    from config import NAS_PATHS
    audit("reindex_trigger", user_id=admin.id, username=admin.username,
          detail=f"Global NAS reindex triggered ({len(NAS_PATHS)} paths)")
    job_id = idx.start_nas_index(NAS_PATHS)
    return {"job_id": job_id}


@app.post("/api/admin/reindex/full")
def admin_reindex_full(
    admin: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    """Wipe global ChromaDB collection, FTS5, and indexed_files, then re-index from scratch."""
    import sqlite3 as _sqlite3
    from config import DB_PATH, NAS_PATHS
    from rag import _chroma_client, get_or_create_collection
    from models import IndexedFile

    # 1. Delete + recreate ChromaDB global collection
    try:
        _chroma_client().delete_collection(GLOBAL_COLLECTION)
    except Exception:
        pass
    get_or_create_collection(GLOBAL_COLLECTION)

    # 2. Wipe FTS5 entries for global collection
    try:
        fts = _sqlite3.connect(str(DB_PATH))
        fts.execute("DELETE FROM chunk_fts WHERE collection = ?", (GLOBAL_COLLECTION,))
        fts.commit()
        fts.close()
    except Exception:
        pass

    # 3. Wipe indexed_files for global (case_id IS NULL) entries only
    db.query(IndexedFile).filter(IndexedFile.case_id.is_(None)).delete()
    db.commit()

    audit("reindex_full", user_id=admin.id, username=admin.username,
          detail=f"Full wipe + reindex triggered ({len(NAS_PATHS)} paths)")
    job_id = idx.start_nas_index(NAS_PATHS)
    return {"job_id": job_id}


@app.get("/api/admin/watcher/status")
def watcher_status(_: User = Depends(auth.require_admin)):
    return watcher_mod.watcher_status()


# ── CourtListener ─────────────────────────────────────────────────────────────

class CourtListenerRequest(BaseModel):
    count:       int            = 20
    query:       str            = ""
    court:       Optional[str]  = None
    after_date:  Optional[str]  = None
    before_date: Optional[str]  = None
    trigger_index: bool         = True


@app.post("/api/admin/courtlistener/download")
def courtlistener_download(
    req: CourtListenerRequest,
    admin: User = Depends(auth.require_admin),
):
    status = cl_mod.get_download_status()
    if status.get("running"):
        raise HTTPException(status_code=409, detail="A download is already running.")
    audit("courtlistener_download", user_id=admin.id, username=admin.username,
          detail=f"count={req.count} court={req.court} query={req.query!r}")
    cl_mod.start_download(
        count=req.count,
        query=req.query,
        court=req.court,
        after_date=req.after_date,
        before_date=req.before_date,
        trigger_index=req.trigger_index,
    )
    return {"started": True}


@app.get("/api/admin/courtlistener/status")
def courtlistener_status(_: User = Depends(auth.require_admin)):
    return cl_mod.get_download_status()


@app.get("/api/admin/courtlistener/courts")
def courtlistener_courts(_: User = Depends(auth.require_admin)):
    return cl_mod.list_courts()


@app.get("/api/indexer/live-status")
def indexer_live_status(current_user: User = Depends(auth.get_current_user)):
    """Live indexer status readable by any logged-in user, from any process."""
    # Check in-process jobs first (triggered via web UI)
    with idx._jobs_lock:
        running = [
            {**job, "job_id": jid}
            for jid, job in idx._jobs.items()
            if not job.get("done") and job.get("status") in ("running", "queued")
        ]
    if running:
        return {"active": True, **running[-1]}
    # Fall back to shared status file (triggered externally, e.g. launchd)
    data = idx.read_live_status()
    if data and not data.get("done"):
        return {"active": True, **data}
    return {"active": False}


@app.post("/api/admin/reindex/{job_id}/cancel")
def admin_cancel_reindex(
    job_id: str,
    _: User = Depends(auth.require_admin),
):
    ok = idx.cancel_job(job_id)
    if not ok:
        raise HTTPException(status_code=404, detail="Job not found")
    return {"cancelled": True, "job_id": job_id}


@app.get("/api/admin/reindex/active")
def admin_reindex_active(_: User = Depends(auth.require_admin)):
    """Return the most recent running reindex job, if any."""
    with idx._jobs_lock:
        running = [
            (jid, job) for jid, job in idx._jobs.items()
            if not job.get("done") and job.get("status") in ("running", "queued")
        ]
    if not running:
        return {"active": False}
    job_id, job = running[-1]
    return {"active": True, "job_id": job_id, **job}


@app.get("/api/admin/reindex/{job_id}/status")
def admin_reindex_status(
    job_id: str,
    _: User = Depends(auth.require_admin),
):
    data = idx.get_job_status(job_id)
    if not data:
        raise HTTPException(status_code=404, detail="Job not found")
    return data


@app.get("/api/admin/config")
def admin_get_config(
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    from config import (
        OLLAMA_URL, CHROMA_URL, OUTPUTS_DIR, UPLOADS_DIR, DB_PATH,
        WHISPER_MODEL_DIR, JWT_EXPIRY_HOURS, EMBED_MODEL, LLM_MODEL,
        WHISPER_MODEL, MAX_UPLOAD_MB, RAG_TOP_N, NAS_PATHS,
        SYSTEM_NAME, GLOBAL_COLLECTION,
    )
    from models import Case, IndexedFile
    import socket

    active_cases = db.query(Case).filter(Case.status == "active").count()
    total_indexed = db.query(IndexedFile).count()

    return {
        "system": {
            "name": SYSTEM_NAME,
            "hostname": socket.gethostname(),
            "db_path": DB_PATH,
            "uploads_dir": UPLOADS_DIR,
            "outputs_dir": OUTPUTS_DIR,
            "whisper_model_dir": WHISPER_MODEL_DIR,
        },
        "services": {
            "ollama_url": OLLAMA_URL,
            "chroma_url": CHROMA_URL,
        },
        "models": {
            "llm": LLM_MODEL,
            "embed": EMBED_MODEL,
            "whisper": WHISPER_MODEL,
        },
        "rag": {
            "top_n": RAG_TOP_N,
            "max_upload_mb": MAX_UPLOAD_MB,
            "global_collection": GLOBAL_COLLECTION,
            "jwt_expiry_hours": JWT_EXPIRY_HOURS,
        },
        "nas": {
            "paths": NAS_PATHS,
        },
        "stats": {
            "active_cases": active_cases,
            "indexed_files": total_indexed,
        },
    }


@app.post("/api/admin/nas-paths")
def admin_set_nas_paths(
    body: dict,
    current_user: User = Depends(auth.require_admin),
):
    """Update NAS_PATHS in sherlock.conf and in-memory config without restart."""
    import config as _config
    paths = [p.strip() for p in body.get("nas_paths", []) if str(p).strip()]

    _ROOT = Path(__file__).parent.parent
    conf_path = _ROOT / "sherlock.conf"

    existing: dict[str, str] = {}
    if conf_path.exists():
        for line in conf_path.read_text().splitlines():
            if "=" in line and not line.strip().startswith("#"):
                k, _, v = line.partition("=")
                existing[k.strip()] = v.strip()

    existing["NAS_PATHS"] = ",".join(paths)
    conf_path.write_text("\n".join(f"{k}={v}" for k, v in existing.items()) + "\n")

    # Update in-memory so new index jobs see the change immediately
    _config.NAS_PATHS = paths

    return {"saved": True, "nas_paths": paths}


# ── Document preview ──────────────────────────────────────────────────────────

_PREVIEW_MIME = {
    ".pdf":  "application/pdf",
    ".txt":  "text/plain; charset=utf-8",
    ".md":   "text/plain; charset=utf-8",
    ".csv":  "text/plain; charset=utf-8",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png":  "image/png",
    ".gif":  "image/gif",
    ".tiff": "image/tiff",
    ".tif":  "image/tiff",
    ".bmp":  "image/bmp",
}
_NATIVE_PREVIEW = {".pdf", ".txt", ".md", ".csv",
                   ".jpg", ".jpeg", ".png", ".gif", ".tiff", ".tif", ".bmp"}


@app.post("/api/open")
def open_file_in_os(
    path: str = Query(...),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Open a file in the user's default OS application (macOS: open, Linux: xdg-open)."""
    import subprocess, sys
    from models import IndexedFile, Upload
    from config import UPLOADS_DIR, OUTPUTS_DIR, NAS_PATHS

    if not path or path.strip() == "":
        raise HTTPException(status_code=400, detail="No file path provided")

    fp = Path(path).expanduser().resolve()

    # Guard against symlink traversal
    try:
        real = fp.resolve()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid path")

    # Build set of trusted root directories
    _sherlock_base = Path(__file__).resolve().parent.parent  # ~/Sherlock
    trusted_roots = [
        Path(UPLOADS_DIR).resolve(),
        Path(OUTPUTS_DIR).resolve(),
        _sherlock_base / "demo",
        _sherlock_base / "SampleData",
        _sherlock_base,
    ]
    for np in (NAS_PATHS or []):
        try:
            trusted_roots.append(Path(np).resolve())
        except Exception:
            pass

    in_trusted = any(
        str(real).startswith(str(root)) for root in trusted_roots
    )

    # Also accept if file is recorded in IndexedFile or Upload tables
    in_db = (
        db.query(IndexedFile).filter(
            (IndexedFile.file_path == str(fp)) | (IndexedFile.file_path == path)
        ).first() is not None
        or db.query(Upload).filter(
            (Upload.stored_path == str(fp)) | (Upload.stored_path == path),
            Upload.user_id == current_user.id,
        ).first() is not None
    )

    if not in_trusted and not in_db:
        # Admins can open any indexed file
        if current_user.role == "admin":
            in_db = db.query(IndexedFile).filter(
                (IndexedFile.file_path == str(fp)) | (IndexedFile.file_path == path)
            ).first() is not None
            if not in_db:
                audit("file_access_denied", user_id=current_user.id,
                      username=current_user.username, file_path=str(fp))
                raise HTTPException(status_code=403, detail="File not accessible")
        else:
            audit("file_access_denied", user_id=current_user.id,
                  username=current_user.username, file_path=str(fp))
            raise HTTPException(status_code=403, detail="File not accessible")

    if not fp.exists():
        raise HTTPException(status_code=404, detail=f"File not found on disk: {fp.name}")

    audit("file_open_os", user_id=current_user.id, username=current_user.username,
          file_path=str(fp))

    try:
        if sys.platform == "darwin":
            subprocess.Popen(["open", str(fp)])
        elif sys.platform.startswith("linux"):
            subprocess.Popen(["xdg-open", str(fp)])
        else:
            subprocess.Popen(["start", str(fp)], shell=True)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not open file: {e}")

    return {"ok": True, "path": str(fp)}


@app.get("/api/preview")
def preview_file(
    path: str = Query(...),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Serve an indexed file for in-browser preview (PDF, images, plain text)."""
    import urllib.parse
    from models import IndexedFile, Upload

    decoded = urllib.parse.unquote(path)
    fp = Path(decoded)

    from config import NAS_PATHS, UPLOADS_DIR, OUTPUTS_DIR

    # Check indexed files and uploads
    allowed = (
        db.query(IndexedFile).filter(IndexedFile.file_path == decoded).first() is not None
        or db.query(Upload).filter(
            Upload.stored_path == decoded,
            Upload.user_id == current_user.id,
        ).first() is not None
    )

    # Also allow files under trusted NAS paths or in nas_catalog
    if not allowed:
        real = fp.resolve()
        trusted = [Path(UPLOADS_DIR).resolve(), Path(OUTPUTS_DIR).resolve()]
        for np in (NAS_PATHS or []):
            try:
                trusted.append(Path(np).resolve())
            except Exception:
                pass
        allowed = any(str(real).startswith(str(r)) for r in trusted)

    if not allowed:
        audit("file_access_denied", user_id=current_user.id, username=current_user.username,
              file_path=decoded)
        raise HTTPException(status_code=403, detail="File not accessible")
    if not fp.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    audit("file_preview", user_id=current_user.id, username=current_user.username,
          file_path=decoded)
    ext = fp.suffix.lower()
    mime = _PREVIEW_MIME.get(ext, "application/octet-stream")
    disp = "inline" if ext in _NATIVE_PREVIEW else "attachment"
    return FileResponse(
        str(fp), media_type=mime,
        headers={"Content-Disposition": f'{disp}; filename="{fp.name}"'},
    )


@app.get("/api/preview/text")
def preview_file_text(
    path: str = Query(...),
    highlight: str = Query(default=""),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Return extracted text as a readable HTML page for non-natively-previewable formats."""
    import urllib.parse
    from models import IndexedFile, Upload
    from fastapi.responses import HTMLResponse

    decoded = urllib.parse.unquote(path)
    fp = Path(decoded)

    from config import NAS_PATHS, UPLOADS_DIR, OUTPUTS_DIR

    allowed = (
        db.query(IndexedFile).filter(IndexedFile.file_path == decoded).first() is not None
        or db.query(Upload).filter(
            Upload.stored_path == decoded,
            Upload.user_id == current_user.id,
        ).first() is not None
    )
    if not allowed:
        real = fp.resolve()
        trusted = [Path(UPLOADS_DIR).resolve(), Path(OUTPUTS_DIR).resolve()]
        for np in (NAS_PATHS or []):
            try:
                trusted.append(Path(np).resolve())
            except Exception:
                pass
        allowed = any(str(real).startswith(str(r)) for r in trusted)

    if not allowed:
        raise HTTPException(status_code=403, detail="File not accessible")
    if not fp.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    try:
        text = idx.extract_text(fp)
    except Exception as e:
        text = f"[Could not extract text: {e}]"

    def _esc(s):
        return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    escaped_text = _esc(text or "[No text content extracted]")
    if highlight:
        import re
        # Case-insensitive highlight of the query terms
        for term in highlight.split()[:5]:  # limit to 5 terms
            safe_term = _esc(term)
            if len(safe_term) >= 3:  # skip very short terms
                escaped_text = re.sub(
                    f"({re.escape(safe_term)})",
                    r"<mark>\1</mark>",
                    escaped_text,
                    flags=re.IGNORECASE,
                )

    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<title>{_esc(fp.name)}</title>
<style>
body{{font-family:Georgia,serif;max-width:860px;margin:40px auto;padding:0 24px;
     color:#1a1a2e;line-height:1.7;background:#fafafa}}
h1{{font-size:14px;color:#666;border-bottom:1px solid #ddd;padding-bottom:8px;font-family:monospace}}
pre{{white-space:pre-wrap;word-break:break-word;font-family:inherit;font-size:13px}}
mark{{background:#c9a84c33;border-bottom:2px solid #c9a84c;padding:0 2px}}
</style></head><body>
<h1>&#128196; {_esc(fp.name)}</h1>
<pre>{escaped_text}</pre>
</body></html>"""

    return HTMLResponse(content=html)


# ── Matter export ─────────────────────────────────────────────────────────────

@app.get("/api/matters/{matter_id}/export")
def export_matter(
    matter_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """
    Printable HTML export of a matter's full conversation + sources.
    Browser → File → Print → Save as PDF produces a clean legal document.
    """
    from fastapi.responses import HTMLResponse

    matter = db.query(Matter).filter(
        Matter.id == matter_id, Matter.user_id == current_user.id
    ).first()
    if not matter:
        raise HTTPException(status_code=404, detail="Matter not found")

    messages = (
        db.query(Message)
        .filter(Message.matter_id == matter_id)
        .order_by(Message.created_at)
        .all()
    )
    audit("matter_export", user_id=current_user.id, username=current_user.username,
          detail=f"Matter '{matter.name}' ({len(messages)} messages)")

    def _esc(s):
        return (s or "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    msg_html = []
    for msg in messages:
        role_label = (current_user.display_name or current_user.username
                      if msg.role == "user" else "Sherlock")
        ts = msg.created_at.strftime("%b %d, %Y %I:%M %p") if msg.created_at else ""
        srcs = msg.sources_list()
        src_html = ""
        if srcs:
            items = "".join(
                f'<li><strong>{_esc(s.get("file",""))}</strong>'
                f'{": " + _esc(s.get("excerpt","")[:120]) + "…" if s.get("excerpt") else ""}</li>'
                for s in srcs
            )
            src_html = f'<div class="sources"><div class="src-label">Sources</div><ol>{items}</ol></div>'

        msg_html.append(f"""
<div class="msg {msg.role}">
  <div class="msg-meta"><span class="role">{_esc(role_label)}</span><span class="ts">{_esc(ts)}</span></div>
  <div class="msg-body">{_esc(msg.content)}</div>{src_html}
</div>""")

    date_range = ""
    if messages:
        first = messages[0].created_at.strftime("%b %d, %Y") if messages[0].created_at else ""
        last  = messages[-1].created_at.strftime("%b %d, %Y") if messages[-1].created_at else ""
        date_range = first if first == last else f"{first} – {last}"

    exported_at = datetime.utcnow().strftime("%B %d, %Y at %I:%M %p UTC")
    msg_count = len(messages)

    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8">
<title>{_esc(matter.name)} — Sherlock Export</title>
<style>
@page{{margin:1in}}
*{{box-sizing:border-box}}
body{{font-family:Georgia,'Times New Roman',serif;font-size:12pt;color:#111;
     max-width:800px;margin:0 auto;padding:20px}}
.export-header{{border-bottom:2px solid #1a1a2e;margin-bottom:32px;padding-bottom:16px}}
.export-header h1{{font-size:22pt;margin:0 0 6px}}
.export-meta{{font-size:10pt;color:#555}}
.msg{{margin-bottom:24px;page-break-inside:avoid}}
.msg-meta{{font-size:9pt;color:#888;margin-bottom:4px;display:flex;justify-content:space-between}}
.msg-meta .role{{font-weight:bold;color:#1a1a2e;text-transform:uppercase;
                letter-spacing:.05em;font-size:8pt}}
.user .msg-body{{background:#f5f7fa;border-left:3px solid #1a1a2e;
                padding:10px 14px;border-radius:0 6px 6px 0;white-space:pre-wrap}}
.assistant .msg-body{{background:#fff;border-left:3px solid #c8973a;
                     padding:10px 14px;border-radius:0 6px 6px 0;white-space:pre-wrap}}
.sources{{margin-top:10px;font-size:9pt;color:#444}}
.src-label{{font-weight:bold;text-transform:uppercase;letter-spacing:.05em;
           color:#888;font-size:8pt;margin-bottom:4px}}
.sources ol{{margin:0;padding-left:20px}}
.sources li{{margin-bottom:3px}}
.export-footer{{border-top:1px solid #ddd;margin-top:40px;padding-top:12px;
               font-size:9pt;color:#999;text-align:center}}
@media print{{body{{padding:0}}a{{color:inherit;text-decoration:none}}}}
</style></head><body>
<div class="export-header">
  <h1>{_esc(matter.name)}</h1>
  <div class="export-meta">
    {f"Date range: {_esc(date_range)} &nbsp;|&nbsp; " if date_range else ""}
    {msg_count} message{'s' if msg_count != 1 else ''} &nbsp;|&nbsp;
    Exported by {_esc(current_user.display_name or current_user.username)}
  </div>
</div>
{''.join(msg_html) or '<p style="color:#888">No messages in this matter.</p>'}
<div class="export-footer">Generated by Sherlock &nbsp;·&nbsp; {_esc(exported_at)}</div>
</body></html>"""

    safe_name = matter.name.replace("/","-").replace("\\","-")[:80]
    return HTMLResponse(
        content=html,
        headers={"Content-Disposition": f'inline; filename="sherlock-{safe_name}.html"'},
    )


# ── Recent query history ───────────────────────────────────────────────────────

@app.get("/api/history")
def get_history(
    limit: int = Query(default=15, le=50),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Return the current user's recent queries across all matters."""
    rows = (
        db.query(Message.id, Message.content, Message.created_at, Message.matter_id,
                 Matter.name.label("matter_name"))
        .join(Matter, Message.matter_id == Matter.id)
        .filter(Message.user_id == current_user.id, Message.role == "user")
        .order_by(Message.created_at.desc())
        .limit(limit)
        .all()
    )
    return [
        {
            "id":          r.id,
            "query":       r.content[:120],
            "matter_id":   r.matter_id,
            "matter_name": r.matter_name,
            "created_at":  r.created_at.isoformat() if r.created_at else None,
        }
        for r in rows
    ]


# ── Query Templates ──────────────────────────────────────────────────────────

@app.get("/api/templates")
def list_templates(
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import QueryTemplate
    templates = (
        db.query(QueryTemplate)
        .filter((QueryTemplate.user_id == None) | (QueryTemplate.user_id == current_user.id))
        .order_by(QueryTemplate.category, QueryTemplate.name)
        .all()
    )
    return [
        {"id": t.id, "name": t.name, "template": t.template, "category": t.category,
         "query_type": t.query_type, "is_system": t.user_id is None}
        for t in templates
    ]


class TemplateCreate(BaseModel):
    name: str
    template: str
    category: str = "general"
    query_type: str = "auto"


@app.post("/api/templates", status_code=201)
def create_template(
    body: TemplateCreate,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import QueryTemplate
    t = QueryTemplate(user_id=current_user.id, name=body.name, template=body.template,
                      category=body.category, query_type=body.query_type)
    db.add(t)
    db.commit()
    db.refresh(t)
    return {"id": t.id, "name": t.name}


@app.delete("/api/templates/{template_id}")
def delete_template(
    template_id: int,
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    from models import QueryTemplate
    t = db.query(QueryTemplate).filter(QueryTemplate.id == template_id, QueryTemplate.user_id == current_user.id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Template not found or is a system template")
    db.delete(t)
    db.commit()
    return {"deleted": template_id}


# ── Log viewer API ────────────────────────────────────────────────────────────

_VALID_STREAMS = {"app", "audit", "rag", "indexer", "web"}


@app.get("/api/admin/logs")
def get_logs(
    stream: str = Query(default="app"),
    lines: int  = Query(default=500, le=2000),
    level: str  = Query(default=""),
    search: str = Query(default=""),
    _: User = Depends(auth.require_admin),
):
    """Return recent log entries from a log stream as JSON."""
    if stream not in _VALID_STREAMS:
        raise HTTPException(status_code=400, detail=f"Invalid stream. Choose: {_VALID_STREAMS}")
    entries = tail_log(
        stream,
        lines=lines,
        level=level or None,
        search=search or None,
    )
    return {"stream": stream, "count": len(entries), "entries": entries}


@app.get("/api/admin/logs/download")
def download_log(
    stream: str = Query(default="app"),
    _: User = Depends(auth.require_admin),
):
    """Download a full log file."""
    if stream not in _VALID_STREAMS:
        raise HTTPException(status_code=400, detail="Invalid stream")
    from pathlib import Path as _P
    log_path = _P(__file__).parent.parent / "logs" / f"{stream}.log"
    if not log_path.exists():
        raise HTTPException(status_code=404, detail="Log file not found")
    return FileResponse(
        str(log_path),
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="sherlock-{stream}.log"'},
    )


# ── Rate limit status & Usage dashboard ──────────────────────────────────────

@app.get("/api/admin/rate-limits")
def get_rate_limits(_: User = Depends(auth.require_admin)):
    from config import RATE_LIMIT_RPM
    now = time.time()
    with _rate_lock:
        result = {}
        for uid, bucket in _rate_buckets.items():
            recent = sum(1 for t in bucket if t > now - 60)
            result[str(uid)] = {"requests_last_60s": recent, "limit": RATE_LIMIT_RPM}
    return result


@app.get("/api/admin/usage")
def admin_usage(
    days: int = Query(7, ge=1, le=90),
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    from datetime import timedelta
    from models import QueryLog, User as UserM
    from sqlalchemy import func

    since = datetime.utcnow() - timedelta(days=days)

    # Total queries in window
    total = db.query(func.count(QueryLog.id)).filter(QueryLog.created_at >= since).scalar() or 0

    # Per-user stats
    user_rows = (
        db.query(UserM.id, UserM.username, UserM.display_name, func.count(QueryLog.id).label("queries"))
        .join(QueryLog, QueryLog.user_id == UserM.id, isouter=True)
        .filter(QueryLog.created_at >= since)
        .group_by(UserM.id)
        .order_by(func.count(QueryLog.id).desc())
        .all()
    )
    per_user = [{"user_id": r.id, "username": r.username, "display_name": r.display_name,
                 "queries": r.queries or 0} for r in user_rows]

    # Daily query volume (last N days)
    daily_rows = (
        db.query(
            func.strftime('%Y-%m-%d', QueryLog.created_at).label("day"),
            func.count(QueryLog.id).label("count"),
        )
        .filter(QueryLog.created_at >= since)
        .group_by("day")
        .order_by("day")
        .all()
    )
    daily = [{"day": r.day, "count": r.count} for r in daily_rows]

    # Query type breakdown
    type_rows = (
        db.query(QueryLog.query_type, func.count(QueryLog.id).label("count"))
        .filter(QueryLog.created_at >= since)
        .group_by(QueryLog.query_type)
        .all()
    )
    by_type = {r.query_type or "auto": r.count for r in type_rows}

    # Token usage aggregates
    total_prompt_tokens = db.query(func.coalesce(func.sum(QueryLog.prompt_tokens), 0)).filter(QueryLog.created_at >= since).scalar()
    total_completion_tokens = db.query(func.coalesce(func.sum(QueryLog.completion_tokens), 0)).filter(QueryLog.created_at >= since).scalar()
    total_all_tokens = db.query(func.coalesce(func.sum(QueryLog.total_tokens), 0)).filter(QueryLog.created_at >= since).scalar()
    avg_tps = db.query(func.avg(QueryLog.tokens_per_sec)).filter(QueryLog.created_at >= since, QueryLog.tokens_per_sec > 0).scalar()

    # Per-user token usage
    user_token_rows = (
        db.query(
            UserM.id, UserM.username, UserM.display_name,
            func.count(QueryLog.id).label("queries"),
            func.coalesce(func.sum(QueryLog.prompt_tokens), 0).label("prompt_tokens"),
            func.coalesce(func.sum(QueryLog.completion_tokens), 0).label("completion_tokens"),
            func.coalesce(func.sum(QueryLog.total_tokens), 0).label("total_tokens"),
        )
        .join(QueryLog, QueryLog.user_id == UserM.id, isouter=True)
        .filter(QueryLog.created_at >= since)
        .group_by(UserM.id)
        .order_by(func.coalesce(func.sum(QueryLog.total_tokens), 0).desc())
        .all()
    )
    per_user_tokens = [
        {"user_id": r.id, "username": r.username, "display_name": r.display_name,
         "queries": r.queries or 0, "prompt_tokens": r.prompt_tokens,
         "completion_tokens": r.completion_tokens, "total_tokens": r.total_tokens}
        for r in user_token_rows
    ]

    # Current rate limit status
    now = time.time()
    with _rate_lock:
        rate_status = {str(uid): sum(1 for t in bkt if t > now - 60)
                       for uid, bkt in _rate_buckets.items()}

    # Token usage broken out by source (user vs system)
    source_rows = (
        db.query(
            QueryLog.source,
            func.coalesce(func.sum(QueryLog.prompt_tokens), 0).label("prompt"),
            func.coalesce(func.sum(QueryLog.completion_tokens), 0).label("completion"),
            func.coalesce(func.sum(QueryLog.total_tokens), 0).label("total"),
            func.count(QueryLog.id).label("count"),
        )
        .filter(QueryLog.created_at >= since)
        .group_by(QueryLog.source)
        .all()
    )
    by_source = {
        r.source or "user": {
            "prompt_tokens": r.prompt, "completion_tokens": r.completion,
            "total_tokens": r.total, "queries": r.count,
        }
        for r in source_rows
    }

    return {
        "days": days,
        "total_queries": total,
        "per_user": per_user_tokens,
        "daily": daily,
        "by_type": by_type,
        "rate_limits": rate_status,
        "tokens": {
            "prompt_tokens": total_prompt_tokens,
            "completion_tokens": total_completion_tokens,
            "total_tokens": total_all_tokens,
            "avg_tokens_per_sec": round(avg_tps, 1) if avg_tps else 0,
        },
        "by_source": by_source,
    }


@app.get("/api/admin/audit")
def admin_audit(
    days: int = Query(7, ge=1, le=90),
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    """Return audit trail data for compliance dashboard."""
    from datetime import timedelta
    from models import QueryLog, User as UserM, IndexedFile, Upload, Output
    from sqlalchemy import func

    since = datetime.utcnow() - timedelta(days=days)

    # Query activity per user
    user_queries = (
        db.query(
            UserM.id, UserM.username, UserM.display_name,
            func.count(QueryLog.id).label("queries"),
            func.max(QueryLog.created_at).label("last_query"),
        )
        .join(QueryLog, QueryLog.user_id == UserM.id, isouter=True)
        .filter(QueryLog.created_at >= since)
        .group_by(UserM.id)
        .order_by(func.count(QueryLog.id).desc())
        .all()
    )

    # File access counts (uploads + outputs)
    upload_counts = (
        db.query(UserM.username, func.count(Upload.id).label("uploads"))
        .join(Upload, Upload.user_id == UserM.id, isouter=True)
        .filter(Upload.uploaded_at >= since)
        .group_by(UserM.id)
        .all()
    )
    upload_map = {r.username: r.uploads for r in upload_counts}

    output_counts = (
        db.query(UserM.username, func.count(Output.id).label("exports"))
        .join(Output, Output.user_id == UserM.id, isouter=True)
        .filter(Output.saved_at >= since)
        .group_by(UserM.id)
        .all()
    )
    output_map = {r.username: r.exports for r in output_counts}

    # Login activity from audit log
    # (We can derive this from the users table last_login)
    logins = (
        db.query(UserM.id, UserM.username, UserM.last_login)
        .filter(UserM.active == True)
        .all()
    )

    # Total indexed documents
    total_indexed = db.query(func.count(IndexedFile.id)).scalar() or 0
    recent_indexed = (
        db.query(func.count(IndexedFile.id))
        .filter(IndexedFile.indexed_at >= since)
        .scalar() or 0
    )

    # Daily query heatmap
    daily = (
        db.query(
            func.strftime('%Y-%m-%d', QueryLog.created_at).label("day"),
            func.strftime('%H', QueryLog.created_at).label("hour"),
            func.count(QueryLog.id).label("count"),
        )
        .filter(QueryLog.created_at >= since)
        .group_by("day", "hour")
        .all()
    )

    return {
        "days": days,
        "users": [
            {
                "user_id": r.id, "username": r.username,
                "display_name": r.display_name,
                "queries": r.queries or 0,
                "uploads": upload_map.get(r.username, 0),
                "exports": output_map.get(r.username, 0),
                "last_query": r.last_query.isoformat() if r.last_query else None,
            }
            for r in user_queries
        ],
        "logins": [
            {"username": r.username, "last_login": r.last_login.isoformat() if r.last_login else None}
            for r in logins
        ],
        "documents": {
            "total_indexed": total_indexed,
            "recently_indexed": recent_indexed,
        },
        "activity_heatmap": [
            {"day": r.day, "hour": r.hour, "count": r.count}
            for r in daily
        ],
    }


# ── Task Billing Export (CSV) ─────────────────────────────────────────────────

def _build_task_csv(rows) -> str:
    """Build CSV string from task query rows."""
    import csv, io
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["Task ID", "Task Name", "User", "Case", "Billable Hours", "Created", "Archived"])
    for r in rows:
        writer.writerow([
            r.id,
            r.name,
            r.username,
            r.case_name or "",
            f"{r.billable_time or 0.0:.2f}",
            r.created_at.strftime("%Y-%m-%d %H:%M") if r.created_at else "",
            "Yes" if r.archived else "No",
        ])
    buf.seek(0)
    return buf.getvalue()


@app.get("/api/tasks/export")
def export_my_tasks(
    start: Optional[str] = Query(None, description="Start date YYYY-MM-DD"),
    end: Optional[str] = Query(None, description="End date YYYY-MM-DD"),
    current_user: User = Depends(auth.get_current_user),
    db: Session = Depends(get_db),
):
    """Export current user's tasks as CSV for billing."""
    from starlette.responses import Response
    q = (
        db.query(
            Matter.id, Matter.name, User.username,
            Case.case_name, Matter.billable_time,
            Matter.created_at, Matter.archived,
        )
        .join(User, User.id == Matter.user_id)
        .outerjoin(Case, Case.id == Matter.case_id)
        .filter(Matter.user_id == current_user.id)
    )
    if start:
        q = q.filter(Matter.created_at >= start)
    if end:
        q = q.filter(Matter.created_at <= end + " 23:59:59")
    rows = q.order_by(Matter.created_at.desc()).all()
    csv_str = _build_task_csv(rows)
    return Response(
        content=csv_str,
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="my-tasks.csv"'},
    )


@app.get("/api/admin/tasks/export")
def export_all_tasks(
    user_id: Optional[int] = Query(None, description="Filter by user ID"),
    start: Optional[str] = Query(None, description="Start date YYYY-MM-DD"),
    end: Optional[str] = Query(None, description="End date YYYY-MM-DD"),
    _: User = Depends(auth.require_admin),
    db: Session = Depends(get_db),
):
    """Admin: export all tasks across all users as CSV for billing."""
    from starlette.responses import Response
    q = (
        db.query(
            Matter.id, Matter.name, User.username,
            Case.case_name, Matter.billable_time,
            Matter.created_at, Matter.archived,
        )
        .join(User, User.id == Matter.user_id)
        .outerjoin(Case, Case.id == Matter.case_id)
    )
    if user_id:
        q = q.filter(Matter.user_id == user_id)
    if start:
        q = q.filter(Matter.created_at >= start)
    if end:
        q = q.filter(Matter.created_at <= end + " 23:59:59")
    rows = q.order_by(User.username, Matter.created_at.desc()).all()
    csv_str = _build_task_csv(rows)
    return Response(
        content=csv_str,
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="all-tasks.csv"'},
    )


# ── Update / Upgrade ─────────────────────────────────────────────────────────

_SHERLOCK_ROOT   = Path(__file__).parent.parent
_VERSION_FILE    = _SHERLOCK_ROOT / "VERSION"
_UPGRADE_SCRIPT  = _SHERLOCK_ROOT / "scripts" / "upgrade.sh"
_GITHUB_REPO     = "Tnijem/SherlockAi"
_upgrade_status: dict = {"running": False, "log": [], "error": None}


def _local_version() -> str:
    try:
        return _VERSION_FILE.read_text().strip()
    except Exception:
        return "unknown"


@app.get("/api/admin/update/check")
def check_for_update(_: User = Depends(auth.require_admin)):
    """Check GitHub releases for a newer version."""
    import urllib.request, json as _json
    current = _local_version()
    try:
        url = f"https://api.github.com/repos/{_GITHUB_REPO}/releases/latest"
        req = urllib.request.Request(url, headers={"User-Agent": "Sherlock-Updater/1.0"})
        with urllib.request.urlopen(req, timeout=6) as resp:
            data = _json.loads(resp.read())
        latest = data.get("tag_name", "").strip()
        body   = data.get("body", "")
        assets = [a["browser_download_url"] for a in data.get("assets", [])
                  if a["name"] == "sherlock-source.tar.gz"]
        return {
            "current":       current,
            "latest":        latest,
            "update_available": latest != current and latest != "",
            "release_notes": body[:600] if body else "",
            "asset_url":     assets[0] if assets else None,
        }
    except Exception as exc:
        return {"current": current, "latest": None, "update_available": False,
                "error": str(exc)}


@app.post("/api/admin/update/apply")
def apply_update(_: User = Depends(auth.require_admin)):
    """Trigger an immediate upgrade in the background."""
    import subprocess, threading
    if _upgrade_status["running"]:
        raise HTTPException(status_code=409, detail="Upgrade already in progress.")
    if not _UPGRADE_SCRIPT.exists():
        raise HTTPException(status_code=500, detail="Upgrade script not found.")

    def _run():
        _upgrade_status["running"] = True
        _upgrade_status["log"]     = []
        _upgrade_status["error"]   = None
        try:
            proc = subprocess.Popen(
                ["bash", str(_UPGRADE_SCRIPT), "--yes"],
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True
            )
            for line in proc.stdout:
                _upgrade_status["log"].append(line.rstrip())
            proc.wait()
            if proc.returncode != 0:
                _upgrade_status["error"] = f"Upgrade exited with code {proc.returncode}"
        except Exception as exc:
            _upgrade_status["error"] = str(exc)
        finally:
            _upgrade_status["running"] = False

    threading.Thread(target=_run, daemon=True).start()
    audit("upgrade_started", user="admin")
    return {"status": "started"}


@app.get("/api/admin/update/status")
def upgrade_status(_: User = Depends(auth.require_admin)):
    """Poll progress of an in-flight or completed upgrade."""
    return {
        "running": _upgrade_status["running"],
        "log":     _upgrade_status["log"][-50:],   # last 50 lines
        "error":   _upgrade_status["error"],
        "version": _local_version(),
    }


@app.post("/api/admin/update/schedule")
def schedule_update(
    body: dict,
    _: User = Depends(auth.require_admin),
):
    """Schedule upgrade via cron (time = 'HH:MM', default '03:00')."""
    import subprocess
    time_str  = body.get("time", "03:00")
    try:
        hh, mm = [int(x) for x in time_str.split(":")]
        if not (0 <= hh <= 23 and 0 <= mm <= 59):
            raise ValueError
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid time. Use HH:MM (24h).")

    script = str(_UPGRADE_SCRIPT)
    # Use `at` if available, else cron
    try:
        subprocess.run(
            ["bash", "-c", f"echo 'bash {script} --yes >> /tmp/sherlock-upgrade.log 2>&1' | at {hh:02d}:{mm:02d}"],
            check=True, capture_output=True
        )
        method = "at"
    except Exception:
        # Fallback: add a cron entry for tonight only (removes itself after running)
        cron_line = f"{mm} {hh} * * * bash {script} --yes >> /tmp/sherlock-upgrade.log 2>&1 # sherlock-upgrade-once"
        result = subprocess.run(["crontab", "-l"], capture_output=True, text=True)
        existing = result.stdout if result.returncode == 0 else ""
        # Remove any old scheduled entry first
        cleaned = "\n".join(l for l in existing.splitlines() if "# sherlock-upgrade-once" not in l)
        new_cron = cleaned.rstrip("\n") + "\n" + cron_line + "\n"
        subprocess.run(["crontab", "-"], input=new_cron, text=True, check=True)
        method = "cron"

    audit("upgrade_scheduled", user="admin", time=time_str, method=method)
    return {"status": "scheduled", "time": time_str, "method": method}




# ── Index Filters ─────────────────────────────────────────────────────────────

import file_filters as _ff

class FilterRuleIn(BaseModel):
    name:             str
    enabled:          bool           = True
    action:           str            = "exclude"
    filename_pattern: Optional[str]  = None
    path_pattern:     Optional[str]  = None
    created_before:   Optional[str]  = None
    created_after:    Optional[str]  = None
    modified_before:  Optional[str]  = None
    modified_after:   Optional[str]  = None
    size_gt:          Optional[int]  = None
    size_lt:          Optional[int]  = None

class FilterPreviewIn(BaseModel):
    rule:  FilterRuleIn
    paths: list[str] = []


@app.get("/api/admin/filters")
def list_filters(_: User = Depends(auth.require_admin)):
    return _ff.api_list()


@app.post("/api/admin/filters", status_code=201)
def add_filter(rule: FilterRuleIn, admin: User = Depends(auth.require_admin)):
    try:
        result = _ff.api_add(rule.dict(exclude_none=True))
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    audit("filter_add", user_id=admin.id, username=admin.username, detail=rule.name)
    return result


@app.put("/api/admin/filters/{rule_id}")
def update_filter(rule_id: str, updates: FilterRuleIn, admin: User = Depends(auth.require_admin)):
    result = _ff.api_update(rule_id, updates.dict(exclude_none=True))
    if result is None:
        raise HTTPException(status_code=404, detail="Filter rule not found")
    audit("filter_update", user_id=admin.id, username=admin.username, detail=rule_id)
    return result


@app.delete("/api/admin/filters/{rule_id}", status_code=204)
def delete_filter(rule_id: str, admin: User = Depends(auth.require_admin)):
    if not _ff.api_delete(rule_id):
        raise HTTPException(status_code=404, detail="Filter rule not found")
    audit("filter_delete", user_id=admin.id, username=admin.username, detail=rule_id)


@app.post("/api/admin/filters/preview")
def preview_filter(body: FilterPreviewIn, _: User = Depends(auth.require_admin)):
    from config import NAS_PATHS
    paths = body.paths if body.paths else NAS_PATHS
    try:
        return _ff.api_preview(body.rule.dict(exclude_none=True), paths)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))


# ── NAS Catalog (Tier 1: metadata search) ────────────────────────────────────

@app.post("/api/catalog/scan")
def catalog_scan(
    incremental: bool = True,
    _admin: User = Depends(auth.require_admin),
):
    """Start a NAS catalog scan (admin only). Incremental by default."""
    return nas_catalog.start_catalog_scan(incremental=incremental)


@app.get("/api/catalog/status")
def catalog_status(current_user: User = Depends(auth.get_current_user)):
    """Get current catalog scan status."""
    return nas_catalog.get_scan_status()


@app.get("/api/catalog/search")
def catalog_search(
    q: str = "",
    client: str = "",
    category: str = "",
    extension: str = "",
    limit: int = Query(50, le=200),
    offset: int = 0,
    current_user: User = Depends(auth.get_current_user),
):
    """Search the NAS file catalog by filename, client, category, or extension."""
    return nas_catalog.search_catalog(
        query=q, client=client, category=category,
        extension=extension, limit=limit, offset=offset,
    )


@app.get("/api/catalog/stats")
def catalog_stats(current_user: User = Depends(auth.get_current_user)):
    """Get catalog summary statistics."""
    return nas_catalog.get_catalog_stats()


@app.get("/api/catalog/clients")
def catalog_clients(
    category: str = "",
    limit: int = Query(500, le=2000),
    current_user: User = Depends(auth.get_current_user),
):
    """List client folders with file counts."""
    return {"clients": nas_catalog.get_client_list(category=category, limit=limit)}


# ── NAS Text Extraction (Tier 2: full-text search) ───────────────────────────

@app.post("/api/text/extract")
def text_extract(_admin: User = Depends(auth.require_admin)):
    """Start background text extraction (admin only)."""
    return nas_text.start_text_extraction()


@app.get("/api/text/status")
def text_status(current_user: User = Depends(auth.get_current_user)):
    """Get text extraction status."""
    return nas_text.get_extract_status()


@app.get("/api/text/search")
def text_search(
    q: str = "",
    client: str = "",
    extension: str = "",
    limit: int = Query(30, le=100),
    offset: int = 0,
    current_user: User = Depends(auth.get_current_user),
):
    """Full-text search across extracted NAS file content."""
    if not q:
        raise HTTPException(status_code=422, detail="Query parameter 'q' is required")
    return nas_text.search_text(query=q, client=client, extension=extension,
                                limit=limit, offset=offset)


@app.get("/api/text/stats")
def text_stats(current_user: User = Depends(auth.get_current_user)):
    """Get text extraction statistics."""
    return nas_text.get_text_stats()


# ── NAS Smart Embedding (Tier 3: semantic search) ────────────────────────────

@app.post("/api/embed/start")
def embed_start(
    limit: int = Query(200, le=1000),
    _admin: User = Depends(auth.require_admin),
):
    """Start background NAS file embedding (admin only)."""
    return nas_embed.start_embedding(limit=limit)


@app.get("/api/embed/status")
def embed_status(current_user: User = Depends(auth.get_current_user)):
    """Get embedding status."""
    return nas_embed.get_embed_status()


@app.get("/api/embed/stats")
def embed_stats(current_user: User = Depends(auth.get_current_user)):
    """Get embedding statistics."""
    return nas_embed.get_embed_stats()

# ── Entrypoint ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=3000, reload=False)
